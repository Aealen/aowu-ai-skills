#!/usr/bin/env python3
"""
build_docx.py —— DOCX 重建（核心模块）

职责：遍历对齐合并后的 blocks，按 type 分发，用 python-docx 重建 DOCX。
这是工作量所在——把 MinerU 的版面结构 + PyMuPDF 的字符样式，
还原成接近"用户手动用 Word 另存的 DOCX"。

验收标准（P0 必须达成）：
  - 标题层级正确：Heading 1/2/3 + OutlineLevel（下游章节分割强依赖）
  - 正文段落完整：无大面积丢失
  - 表格结构还原：行列基本对齐（复杂合并单元格留 TODO）
  - 字号字体基本还原

依赖：python-docx

用法（被 pdf2docx.py 调用）:
    from build_docx import build_docx
    build_docx(merged_data, images_dir, "output/结果.docx")

⚠️ 骨架阶段实现边界：
  - title（含 OutlineLevel）：完整实现 ✅
  - text：完整实现 ✅
  - table：HTML 解析框架搭好，rowspan/colspan 合并留 TODO ⏳
  - image：基本实现（按顺序插入）✅
  - list：降级为普通段落 ✅
"""

from __future__ import annotations

import html
import re
import sys
from pathlib import Path
from typing import Any


# ═══════════════════════════════════════════════════════════════
#  文本提取辅助
# ═══════════════════════════════════════════════════════════════

def _extract_text(block: dict[str, Any]) -> str:
    """
    从 MinerU block 提取纯文本。
    兼容多种结构：block.lines.spans.content / block.content
    """
    # 优先从 lines.spans 拼接
    parts = []
    for line in block.get("lines", []):
        for span in line.get("spans", []):
            content = span.get("content") or span.get("text") or ""
            parts.append(content)
    if parts:
        return "".join(parts)

    # 回退：顶层 content 字段
    return block.get("content", "")


def _get_dominant_style(block: dict[str, Any]) -> dict[str, Any] | None:
    """
    取 block 内占比最大的样式（用于 title 等需要整体设样式的块）。
    优先用 block._style（align.py 兜底贴的），否则统计内部 span。
    """
    # align.py 兜底贴的 block 级样式
    if block.get("_style"):
        return block["_style"]

    # 统计 lines.spans 内的样式
    style_counter: dict[str, int] = {}
    style_map: dict[str, dict] = {}
    for line in block.get("lines", []):
        for span in line.get("spans", []):
            st = span.get("_style")
            if not st:
                continue
            # 用 size 作为指纹（标题通常统一字号）
            key = str(st.get("size", 0))
            style_counter[key] = style_counter.get(key, 0) + 1
            style_map[key] = st

    if not style_counter:
        return None
    dominant_key = max(style_counter, key=style_counter.get)
    return style_map.get(dominant_key)


# ═══════════════════════════════════════════════════════════════
#  OOXML 辅助：OutlineLevel（P0 关键）
# ═══════════════════════════════════════════════════════════════

def _set_outline_level(paragraph, level: int) -> None:
    """
    显式设置段落的 w:outlineLvl。

    ⚠️ P0 关键：下游 splitDocumentByChapters 的 getOutlineLevel()
    强依赖此属性（ChapterSplitServiceImpl.java:268/354/1485-1497）。
    add_heading() 会设 Heading 样式，但显式设 outlineLvl 是双保险。

    依据：python-docx #746，OOXML 规范 w:outlineLvl。
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = paragraph._p.get_or_add_pPr()
    # 先移除已存在的 outlineLvl（避免重复）
    for existing in pPr.findall(qn("w:outlineLvl")):
        pPr.remove(existing)
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level))
    pPr.append(outline)


# ═══════════════════════════════════════════════════════════════
#  样式应用
# ═══════════════════════════════════════════════════════════════


def _get_dominant_size(block: dict[str, Any] | None) -> float:
    """
    从 block 内 span 样式取占比最大的字号。
    用于动态计算行距（字号 × 1.4）。
    """
    if not block:
        return 12.0
    size_counter: dict[float, int] = {}
    for line in block.get("lines", []):
        for sp in line.get("spans", []):
            st = sp.get("_style")
            if st and st.get("size"):
                sz = st["size"]
                size_counter[sz] = size_counter.get(sz, 0) + 1
    if not size_counter:
        return 12.0
    return max(size_counter, key=size_counter.get)


def _detect_page_layout(pdf_info: list, margins: dict | None = None) -> None:
    """
    从每页的 block bbox 数据中检测页面布局参数，注入到每个 block 中。

    检测内容：
      - _page_x0: 正文左边距（优先用 PyMuPDF 精确检测的 margins，回退到 bbox 统计）
      - _page_width: 页宽（page_size[0]）
      - _page_center: 页面水平中心
    """
    from collections import Counter

    # 文档级左边距：优先用 PyMuPDF 精确检测的值，回退到 bbox 统计
    if margins and "left" in margins:
        doc_x0 = margins["left"]
    else:
        all_text_x0s = []
        for page in pdf_info:
            blocks = page.get("para_blocks") or page.get("blocks") or []
            for block in blocks:
                bbox = block.get("bbox")
                if not (bbox and len(bbox) >= 4):
                    continue
                btype = (block.get("type") or block.get("block_type") or "").lower()
                if btype not in ("text", "paragraph", "list"):
                    continue
                raw_text = _extract_text(block).strip()
                if len(raw_text) < 15:
                    continue
                all_text_x0s.append(bbox[0])
        doc_x0 = 82
        if all_text_x0s:
            all_text_x0s.sort()
            doc_x0 = round(all_text_x0s[len(all_text_x0s) // 4] / 3) * 3

    for page in pdf_info:
        page_size = page.get("page_size", [595, 842])
        page_width = page_size[0] if len(page_size) > 0 else 595
        page_center = page_width / 2

        blocks = page.get("para_blocks") or page.get("blocks") or []
        for block in blocks:
            block["_page_x0"] = doc_x0
            block["_page_width"] = page_width
            block["_page_center"] = page_center


def _apply_paragraph_format(para, block: dict[str, Any] | None = None) -> None:
    """
    根据 block 的 bbox 位置推断段落格式。
    所有参数均从 PDF 实际数据提取（_page_x0 / _page_width / _page_center），
    不硬编码任何值。
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    pf = para.paragraph_format
    pf.space_before = Pt(0)

    line_height = (
        (block.get("_line_height") if block else None) or
        (block.get("_doc_line_height") if block else None)
    )
    if line_height:
        pf.line_spacing = Pt(line_height)

    if not block:
        return

    page_x0 = block.get("_page_x0", 82)
    page_width = block.get("_page_width", 595)
    page_center = block.get("_page_center", page_width / 2)

    bbox = block.get("bbox", [])
    lines = block.get("lines", [])

    if not bbox or not lines:
        return

    first_spans = lines[0].get("spans", [])
    if not first_spans:
        return

    first_x0 = first_spans[0].get("bbox", [0])[0]

    # 居中判断：检查各行中心是否接近页面中心
    # 多行 block 中，第一行可能占满全宽（拉宽 block bbox），但后续行是居中的
    # 因此看"多数行的中心是否接近页面中心"，而非 block 整体 bbox
    # 注意：全宽行（占满内容区）的中心天然接近页面中心，但不是居中
    # 只有"行宽 < 内容区宽度且中心接近页面中心"的行才算居中行
    # 额外排除：左对齐文本恰好中心接近页面中心的情况——
    # 居中行的左侧应有明显留白（x0 远离左边距），左对齐的 x0 接近左边距
    content_width = page_width - page_x0 * 2
    centered_lines = 0
    total_lines = 0
    for line in lines:
        spans = line.get("spans", [])
        if not spans:
            continue
        line_bbox = line.get("bbox") or spans[0].get("bbox", [])
        if len(line_bbox) >= 4:
            line_center = (line_bbox[0] + line_bbox[2]) / 2
            line_width = line_bbox[2] - line_bbox[0]
            # 居中行的期望 x0 = 页面中心 - 行宽/2
            expected_x0 = page_center - line_width / 2
            total_lines += 1
            # 居中行：中心接近页面中心 + 行宽未占满 + 实际 x0 接近期望 x0
            if (abs(line_center - page_center) < page_width * 0.05 and
                    line_width < content_width * 0.90 and
                    abs(line_bbox[0] - expected_x0) < page_width * 0.05):
                centered_lines += 1

    if total_lines > 0 and centered_lines >= total_lines / 2:
        is_centered = True
    elif bbox and len(bbox) >= 4:
        # 单行 block 回退到 bbox 对称性判断
        left_gap = bbox[0] - page_x0
        right_gap = page_width - page_x0 - bbox[2]
        gap_diff = abs(left_gap - right_gap)
        is_centered = (gap_diff < page_width * 0.05 and
                       left_gap > page_width * 0.02 and
                       right_gap > page_width * 0.02)
    else:
        is_centered = False

    if is_centered:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Pt(0)
    else:
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # 缩进：分两层
        # left_indent = 续行 x0 相对左边距的偏移
        # first_line_indent = 首行 x0 相对续行 x0 的额外偏移（悬挂缩进）
        # 续行参考用所有 line bbox（MinerU 原始）的最小 x0，不用 span bbox
        # （字体拆分会产生估算 bbox 的 span，污染统计）。
        # 单行 block 仅有首行无续行，left_indent=0 first_line=first_x0-page_x0。
        min_line_x0 = first_x0
        line_x0s = []
        for line in lines:
            lbbox = line.get("bbox", [])
            if lbbox and len(lbbox) >= 4:
                lx0 = round(lbbox[0])
                line_x0s.append(lx0)
                if lx0 < min_line_x0:
                    min_line_x0 = lx0
        if line_x0s:
            # 续行参考：用最小 x0（代表所有续行中最靠左的位置）
            ref_x0 = min(line_x0s)
        else:
            ref_x0 = first_x0
        # left_indent：续行的起始位置
        left_indent_pt = max(0, ref_x0 - page_x0)
        if left_indent_pt > 6:
            pf.left_indent = Pt(round(left_indent_pt))
        # first_line_indent：首行相对续行的额外缩进
        first_indent_pt = max(0, first_x0 - ref_x0)
        if first_indent_pt > 6:
            pf.first_line_indent = Pt(round(first_indent_pt))
        else:
            pf.first_line_indent = Pt(0)

    # 段后间距：从同页下一个 block 的 Y 坐标差获取（与标题格式一致）
    gap_after = block.get("_gap_after")
    if gap_after is not None and gap_after > 0:
        pf.space_after = Pt(round(gap_after))
    else:
        pf.space_after = Pt(0)


def _apply_style(run, style: dict[str, Any] | None = None) -> None:
    """
    把 PyMuPDF 样式应用到 python-docx run。
    style 格式：{"font": str, "size": float, "bold": bool, "italic": bool, "color": int}

    ⚠️ CID 字体加粗检测：PyMuPDF 的 font flags 对某些中文字体（如 SimHei）
    无法准确报告加粗状态（字形轮廓被加粗但 flags 未变）。这里用字体名做补充判断。
    """
    if not style:
        return

    from docx.shared import Pt, RGBColor

    size = style.get("size")
    if size:
        run.font.size = Pt(size)

    # 加粗：完全依赖提取层的判断（flags + 像素密度检测），不再做字体名猜测
    if style.get("bold"):
        run.font.bold = True

    if style.get("italic"):
        run.font.italic = True

    if style.get("underline"):
        run.font.underline = True

    color = style.get("color")
    if color is not None:
        # sRGB 整数 → RGBColor
        # 0=黑，16777215=白
        r = (color >> 16) & 0xFF
        g = (color >> 8) & 0xFF
        b = color & 0xFF
        # 纯黑（0）不设，用默认即可，避免不必要的 XML
        if color != 0:
            run.font.color.rgb = RGBColor(r, g, b)

    font = style.get("font")
    if font:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        font_name = font  # 保留 PDF 原始字体名，由 Word/LibreOffice 自动匹配系统字体
        run.font.name = font_name
        # 中文字体需设 EastAsia 属性
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), font_name)


# ═══════════════════════════════════════════════════════════════
#  各 block 类型的重建函数
# ═══════════════════════════════════════════════════════════════

def _build_title(doc, block: dict[str, Any]) -> None:
    """
    标题：Heading 样式 + OutlineLevel + 字号字体。
    P0 验收项：层级必须正确。
    保留多行结构（封面标题常有多个独立行，各有不同字号）。
    """
    text = _extract_text(block).strip()
    if not text:
        return

    # 标题层级：优先用 PDF 大纲（_toc_level），这是作者在 Word 里设的权威层级
    # 无大纲数据时回退到 MinerU text_level
    level = block.get("_toc_level")
    is_structural = level is not None  # 是否是真·结构标题（在 PDF 大纲中）
    if level is None:
        level = block.get("text_level") or block.get("level") or 1
    try:
        level = int(level)
    except (TypeError, ValueError):
        level = 1
    level = min(max(level, 1), 9)

    # 建段落：结构标题用 Heading 样式 + TOC 绑定，视觉标题用普通段落
    # （视觉标题保留加粗/字号等视觉强调，但不污染 DOCX 目录）
    if is_structural:
        heading_style_name = f"Heading {level}"
        try:
            para = doc.add_paragraph(style=heading_style_name)
        except KeyError:
            para = doc.add_paragraph()
    else:
        para = doc.add_paragraph()

    # 段落格式：所有标题都用 _apply_title_format（LEFT/CENTER，不用 JUSTIFY）
    # _apply_title_format 已支持从 bbox 检测缩进，视觉标题也能正确缩进
    _apply_title_format(para, level, block)

    # 逐行建 run，保留多行结构（每行末尾加换行）
    lines = block.get("lines", [])
    if lines:
        for li, line in enumerate(lines):
            for sp in line.get("spans", []):
                content = sp.get("content") or sp.get("text") or ""
                if not content:
                    continue
                run = para.add_run(content)
                _apply_style(run, sp.get("_style"))
            # 行之间加换行（最后一行不加）
            if li < len(lines) - 1:
                run = para.add_run()
                run.add_break()
    else:
        run = para.add_run(text)
        _apply_style(run, block.get("_style"))

    # 只有结构标题才设 OutlineLevel（绑定到 DOCX TOC）
    if is_structural:
        _set_outline_level(para, level)

    # 强制标题颜色为黑色（覆盖 Heading 样式默认蓝色）
    _force_heading_color_black(para)


def _apply_title_format(para, level: int, block: dict | None = None) -> None:
    """
    标题段落格式。所有参数从 PDF 实际数据提取。
    
    对齐：根据 bbox 位置判断（居中 / 左对齐），不再按 level 写死
    缩进：根据 bbox 实际 x0 计算，不再写死为 0
    行距：优先用 PDF 测量值，无数据时单倍
    段后间距：根据同页下一个 block 的 Y 间距计算
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

    pf = para.paragraph_format
    pf.space_before = Pt(0)

    # 行距
    line_height = (
        (block.get("_line_height") if block else None) or
        (block.get("_doc_line_height") if block else None)
    )
    if line_height:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_height)
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    if not block:
        return

    # 对齐和缩进：从 bbox 位置推断
    page_x0 = block.get("_page_x0", 82)
    page_width = block.get("_page_width", 595)
    page_center = block.get("_page_center", page_width / 2)

    bbox = block.get("bbox", [])
    lines = block.get("lines", [])
    first_x0 = page_x0

    if bbox and lines:
        first_spans = lines[0].get("spans", [])
        if first_spans:
            first_x0 = first_spans[0].get("bbox", [bbox[0]])[0]

    # 居中检测：看多数行的中心是否接近页面中心（排除全宽行）
    content_width = page_width - page_x0 * 2
    centered_lines = 0
    total_lines = 0
    for line in lines:
        spans = line.get("spans", [])
        if not spans:
            continue
        line_bbox = line.get("bbox") or spans[0].get("bbox", [])
        if len(line_bbox) >= 4:
            line_center = (line_bbox[0] + line_bbox[2]) / 2
            line_width = line_bbox[2] - line_bbox[0]
            expected_x0 = page_center - line_width / 2
            total_lines += 1
            if (abs(line_center - page_center) < page_width * 0.05 and
                    line_width < content_width * 0.90 and
                    abs(line_bbox[0] - expected_x0) < page_width * 0.05):
                centered_lines += 1

    if total_lines > 0 and centered_lines >= total_lines / 2:
        is_centered = True
    elif bbox and len(bbox) >= 4:
        left_gap = bbox[0] - page_x0
        right_gap = page_width - page_x0 - bbox[2]
        gap_diff = abs(left_gap - right_gap)
        is_centered = (gap_diff < page_width * 0.05 and
                       left_gap > page_width * 0.02 and
                       right_gap > page_width * 0.02)
    else:
        is_centered = False

    if is_centered:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 标题缩进：用 left_indent（统一所有行），取首行 x0 差值
        # 但如果首行文字已接近内容区宽度（>85%），不设缩进避免溢出换行
        indent_pt = max(0, first_x0 - page_x0)
        content_width = page_width - page_x0 * 2  # 粗估内容宽度
        if indent_pt > 6 and bbox and len(bbox) >= 4:
            block_width = bbox[2] - bbox[0]
            if block_width < content_width * 0.85:
                pf.left_indent = Pt(round(indent_pt))
    pf.first_line_indent = Pt(0)

    # 段后间距：从同页下一个 block 的 Y 坐标差获取，而非硬编码
    gap_after = block.get("_gap_after")
    if gap_after is not None and gap_after > 0:
        pf.space_after = Pt(round(gap_after))
    else:
        pf.space_after = Pt(0)


def _force_heading_color_black(para) -> None:
    """
    强制标题段落所有 run 的颜色为黑色。
    python-docx 的 Heading 样式默认是蓝色（accent1: 365F91/4F81BD），
    源 PDF 的标题都是纯黑（color=0），必须覆盖。
    注意：_apply_style 对 color=0 跳过设置，所以需要单独强制。
    """
    from docx.shared import RGBColor
    for run in para.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def _build_text(doc, block: dict[str, Any]) -> None:
    """
    正文段落。
    按 MinerU 的 lines.spans 粒度设置 run 样式（同段不同样式）。
    """
    text = _extract_text(block)
    if not text.strip():
        return

    para = doc.add_paragraph()

    # 设置段落格式（根据 block bbox 推断对齐/缩进）
    _apply_paragraph_format(para, block)

    # 按 lines.spans 粒度逐个建 run，保留行内混排
    has_lines = bool(block.get("lines"))
    if has_lines:
        for line in block.get("lines", []):
            for mspan in line.get("spans", []):
                content = mspan.get("content") or mspan.get("text") or ""
                if not content:
                    continue
                run = para.add_run(content)
                _apply_style(run, mspan.get("_style"))
    else:
        # 无 lines 结构，整块作为一个 run
        run = para.add_run(text)
        _apply_style(run, block.get("_style"))


def _find_matching_pymupdf_table(
    fitz_doc, page_idx: int, table_bbox: list[float],
    iou_threshold: float = 0.3,
):
    """
    在指定页面用 IoU 匹配 MinerU table bbox 对应的 PyMuPDF 表格。
    返回匹配的 fitz.Table 对象，匹配失败返回 None。
    """
    page = fitz_doc[page_idx]
    tabs = page.find_tables()
    target_x0, target_y0, target_x1, target_y1 = table_bbox

    best_table = None
    best_iou = 0.0
    for t in tabs.tables:
        tx0, ty0, tx1, ty1 = t.bbox
        ix0 = max(target_x0, tx0)
        iy0 = max(target_y0, ty0)
        ix1 = min(target_x1, tx1)
        iy1 = min(target_y1, ty1)
        if ix0 < ix1 and iy0 < iy1:
            inter = (ix1 - ix0) * (iy1 - iy0)
            area_target = (target_x1 - target_x0) * (target_y1 - target_y0)
            area_detected = (tx1 - tx0) * (ty1 - ty0)
            union = area_target + area_detected - inter
            iou = inter / union if union > 0 else 0
            if iou > best_iou:
                best_iou = iou
                best_table = t

    if best_table is None or best_iou < iou_threshold:
        return None
    return best_table


def _extract_table_col_widths_from_table(pymupdf_table) -> list[int] | None:
    """
    从 PyMuPDF Table 对象提取列宽（pt）。
    收集所有 cell 的 x 边界，排序去重 → 列宽。
    """
    xs: set[int] = set()
    for cell in pymupdf_table.cells:
        if cell:
            xs.add(round(cell[0]))
            xs.add(round(cell[2]))
    xs_sorted = sorted(xs)
    if len(xs_sorted) < 2:
        return None
    return [xs_sorted[i + 1] - xs_sorted[i] for i in range(len(xs_sorted) - 1)]


def _extract_table_cell_texts(pymupdf_table) -> list[list[str]] | None:
    """
    从 PyMuPDF Table 对象提取单元格文本矩阵（保留换行 \\n）。

    PyMuPDF 的 extract() 返回 list[list[str]]，每个 cell 的文本
    保留了原文的换行（\\n）。这是恢复表格内换行格式的关键数据源——
    MinerU 的 HTML 把单元格内多行内容合并成了无换行的扁平字符串。

    返回 None 表示提取失败。
    """
    try:
        return pymupdf_table.extract()
    except Exception:
        return None


def _extract_table_cell_styles(
    pymupdf_table, page,
) -> list[list[list[dict[str, Any]]]]:
    """
    从 PDF dict 提取表格每个 cell 内所有 span 的样式（字体/字号/粗体/颜色）。

    返回 cell_styles[r][c] = [span_data, ...]，与 _pymupdf_cells 尺寸平行。
    每个 span_data: {"text", "font", "size", "bold", "italic", "color"}。

    用于在 _fill_table_cells 中按 cell 实际 PDF 样式构建多 run 内容，
    而非全表共用 block._style 一种样式。
    """
    d = page.get_text("dict")
    cell_styles: list[list[list[dict[str, Any]]]] = []

    for row in pymupdf_table.rows:
        row_cells: list[list[dict[str, Any]]] = []
        for cell_bbox in row.cells:
            if not cell_bbox:
                row_cells.append([])
                continue
            cx0, cy0, cx1, cy1 = cell_bbox
            spans: list[dict[str, Any]] = []
            for block in d.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for sp in line.get("spans", []):
                        if not sp["text"].strip():
                            continue
                        bx = sp["bbox"]
                        if not (cx0 - 3 <= bx[0] and bx[2] <= cx1 + 3 and
                                cy0 - 3 <= bx[1] and bx[3] <= cy1 + 3):
                            continue
                        flags = sp.get("flags", 0)
                        spans.append({
                            "text": sp["text"],
                            "font": sp.get("font", ""),
                            "size": sp.get("size", 12),
                            "bold": bool(flags & 16),
                            "italic": bool(flags & 2),
                            "color": sp.get("color", 0),
                            "_y": bx[1],
                            "_x": bx[0],
                        })
            # 排序：先按 y 排序，再按视觉行分组（y 差 < 5pt 归为同一行），
            # 同一视觉行内按 x 坐标排序（从左到右）。
            # 数字/英文和中文的 baseline 有 2-3pt 微小差异，不分组会导致顺序错乱
            # （数字 baseline 略高，排到同行的中文前面）。
            spans.sort(key=lambda s: s["_y"])
            # 分组：相邻 span y 差 < 5pt 归为同一视觉行
            grouped: list[list[dict[str, Any]]] = []
            for sp in spans:
                if grouped and abs(sp["_y"] - grouped[-1][-1]["_y"]) < 5:
                    grouped[-1].append(sp)
                else:
                    grouped.append([sp])
            # 每组内按 x 排序，然后展平
            spans = [sp for group in grouped for sp in sorted(group, key=lambda s: s["_x"])]
            row_cells.append(spans)
        cell_styles.append(row_cells)

    return cell_styles


def _extract_table_row_geometry(
    pymupdf_table, page,
) -> tuple[list[float], float | None, list[float | None]]:
    """
    从 PyMuPDF Table 对象提取每行高度和表格内文字行高。

    参数：
        pymupdf_table: PyMuPDF Table 对象（find_tables 返回）
        page: 所属的 fitz.Page 对象

    返回 (row_heights, text_line_height, row_line_heights)：
      - row_heights: 每行的高度(pt)，用于设置 DOCX trHeight
      - text_line_height: 表格内文字的行间距中位数(pt)，全局兜底
      - row_line_heights: 每行的文字行间距(pt)，与 row_heights 平行。
        用于按行精确设置 cell 段落行距。单行内容时为 None。
    """
    row_heights: list[float] = []
    row_line_heights: list[float | None] = []
    all_text_gaps: list[float] = []
    d = page.get_text("dict")
    for row in pymupdf_table.rows:
        row_h = row.bbox[3] - row.bbox[1]
        # 2% 微缩：补偿 DOCX 渲染与 PDF 的微小差异，控制总页数
        scaled_h = row_h * 0.98
        row_heights.append(round(scaled_h, 1))

        ry0, ry1 = row.bbox[1], row.bbox[3]

        # 按列分组收集 y0s：同一列内的 y0 差值才是行间距，
        # 不同列之间的 y0 差（如 col0 "1.6" 和 col1 "是否允许..."）应忽略。
        col_y0s: dict[int, list[float]] = {}
        for ci, cell_bbox in enumerate(row.cells):
            if not cell_bbox:
                continue
            cx0, _, cx1, _ = cell_bbox
            y0s = []
            for block in d.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    spans = [sp for sp in line.get("spans", []) if sp["text"].strip()]
                    if not spans:
                        continue
                    ly0 = line["bbox"][1]
                    lx0 = line["bbox"][0]
                    if (ry0 - 3 <= ly0 <= ry1 + 3 and
                            cx0 - 3 <= lx0 <= cx1 + 3):
                        y0s.append(round(ly0, 1))
            # 去重：同一行内多个 span（如"年""月""日"各自独立）有相同 y0，
            # 应归为同一行，不能重复计数（否则 max_text_lines 虚高→行距偏小）
            y0s = sorted(set(y0s))
            col_y0s[ci] = y0s

        # 该行的最大文字行数（取所有列中行数最多的）
        max_text_lines = max((len(y0s) for y0s in col_y0s.values()), default=0)

        # 收集同列内连续 y0 差值作为行间距参考
        row_gaps = []
        for y0s in col_y0s.values():
            for i in range(1, len(y0s)):
                gap = y0s[i] - y0s[i - 1]
                if 5 < gap < 40:
                    row_gaps.append(round(gap, 1))
                    all_text_gaps.append(round(gap, 1))

        # 行间距：取 min(文字间距中位数, 行高/(文字行数+0.3))
        # 分母 +0.3 给 cell 上下内边距留余量，避免 exact 模式下
        # N行×line_spacing 恰好等于 trHeight 时最后一行被裁剪
        gap_median = 0.0
        if row_gaps:
            row_gaps.sort()
            gap_median = row_gaps[len(row_gaps) // 2]

        if max_text_lines >= 1:
            # line_spacing = scaled_h / (行数+0.5)，+0.5 给 cell 上下内边距留余量，
            # 避免 exact 模式下 N行×line_spacing 恰好等于 trHeight 时最后一行被裁剪
            by_height = round(scaled_h / (max_text_lines + 0.5), 1)
            if gap_median > 0:
                row_line_heights.append(min(gap_median, by_height))
            else:
                row_line_heights.append(by_height)
        else:
            row_line_heights.append(None)

    text_lh = None
    if all_text_gaps:
        all_text_gaps.sort()
        text_lh = all_text_gaps[len(all_text_gaps) // 2]

    return row_heights, text_lh, row_line_heights


def _measure_block_line_height(
    bbox: list[float], n_lines: int = 1,
) -> float | None:
    """
    测量 block 的实际行高，用于 DOCX 的 EXACTLY line_spacing。

    方法：用 block bbox 的实际高度除以 MinerU 识别的行数。
    这样 line_spacing × 行数 = block 的实际垂直占用，
    精确还原 PDF 中该 block 所占的高度。

    原理验证：
      多行 block: bbox高63pt / 3行 = 21pt → DOCX: 21pt×3 = 63pt = 原始高度 ✓
      单行 block: bbox高14pt / 1行 = 14pt → DOCX: 14pt = 字符高度 ✓
      配合 space_after = gap_after，总占用 = bbox高 + gap = 原PDF段间距 ✓

    返回 None 表示 bbox 无效（调用方回退到默认行距）。
    """
    if not bbox or len(bbox) < 4 or n_lines < 1:
        return None
    block_h = bbox[3] - bbox[1]
    if block_h <= 0:
        return None
    return round(block_h / n_lines, 1)


def _split_crosspage_blocks(pdf_info: list) -> int:
    """
    拆分跨页 text block：按 PDF 页边界切成"本页部分"+"续页部分"。

    MinerU 把跨页段落的全部行（含下一页续行）塞进同一个 block 的 lines 数组，
    但 block 的 bbox 只覆盖本页部分。本函数把它拆成两个独立 block，
    使每个 block 的 bbox 准确覆盖自己的行，从而 bbox高/行数 能算出正确行高；
    配合主循环的分页符实现 PDF 分页复刻。

    数据流：
      1. 检测跨页 block（lines y 跨度 > bbox 高 × 1.5）
      2. 本页行 = y0 落在 bbox 范围内（容差 5pt）的 line
         续页行 = 其余 line（y0 是下一页坐标，突然变小≈80）
      3. 原 block 只保留本页行（lines 截断）
      4. 下一页的空占位 block#0（lines_deleted:true, lines:[]）
         被替换为续页 block：塞入续页行，继承占位 block 的 bbox（已精确覆盖续行区域），
         标记 _crosspage_continuation=True（主循环据此插分页符）

    返回拆分的 block 数量（用于日志）。
    """
    split_count = 0
    for pi, page in enumerate(pdf_info):
        if not isinstance(page, dict):
            continue
        blocks = page.get("para_blocks") or page.get("blocks") or []
        for bi, block in enumerate(blocks):
            btype = (block.get("type") or block.get("block_type") or "").lower()
            if btype not in ("text", "paragraph", "list"):
                continue  # 只拆文本类 block
            bbox = block.get("bbox") or []
            lines = block.get("lines") or []
            if not bbox or len(bbox) < 4 or not lines:
                continue

            # 收集所有 line 的 y 坐标
            line_ys = []
            for ln in lines:
                lbbox = ln.get("bbox")
                if lbbox and len(lbbox) >= 4:
                    line_ys.append((lbbox[1], lbbox[3]))
            if not line_ys:
                continue

            # 跨页检测（判据：lines y 跨度 > bbox 高 × 1.5）
            # 正常 block：lines_span ≈ block_h；跨页 block：lines_span ≈ 满页高（≈660pt）
            block_h = bbox[3] - bbox[1]
            if block_h <= 0:
                continue
            lines_span = max(y1 for _, y1 in line_ys) - min(y0 for y0, _ in line_ys)
            if lines_span <= block_h * 1.5:
                continue  # 非跨页，跳过

            # 跨页 block：按 y0 分本页行 / 续页行
            on_page_lines = [ln for ln in lines
                             if ln.get("bbox") and len(ln["bbox"]) >= 4
                             and ln["bbox"][1] >= bbox[1] - 5]
            cont_lines = [ln for ln in lines if ln not in on_page_lines]
            if not on_page_lines or not cont_lines:
                continue  # 无法拆分，跳过

            # 1. 原 block 只保留本页行
            block["lines"] = on_page_lines

            # 2. 找下一页的空占位 block#0，替换为续页 block
            next_page = pdf_info[pi + 1] if pi + 1 < len(pdf_info) else None
            if not isinstance(next_page, dict):
                continue
            next_blocks = next_page.get("para_blocks") or next_page.get("blocks") or []

            # 空占位特征：type=text/paragraph, lines=[], lines_deleted=true
            placeholder = None
            if next_blocks:
                ph = next_blocks[0]
                ph_type = (ph.get("type") or ph.get("block_type") or "").lower()
                ph_lines = ph.get("lines") or []
                if (ph_type in ("text", "paragraph")
                        and len(ph_lines) == 0
                        and ph.get("lines_deleted")):
                    placeholder = ph

            if placeholder:
                # 用续页行填充占位 block（bbox 已精确覆盖续行区域，保留）
                placeholder["lines"] = cont_lines
                placeholder["type"] = block.get("type") or "text"
                if "block_type" in placeholder:
                    placeholder["block_type"] = block.get("block_type") or "text"
                placeholder["_crosspage_continuation"] = True
            else:
                # 无占位 block 的兜底：在下一页 blocks 开头插入续页 block
                cont_bbox = list(bbox)  # 回退用原 bbox
                cont_block = {
                    "type": block.get("type") or "text",
                    "bbox": cont_bbox,
                    "lines": cont_lines,
                    "_crosspage_continuation": True,
                }
                if next_page.get("para_blocks") is not None:
                    next_page["para_blocks"].insert(0, cont_block)
                if next_page.get("blocks") is not None:
                    next_page["blocks"].insert(0, cont_block)

            split_count += 1

    return split_count


def _match_blocks_with_toc(fitz_doc, pdf_info: list) -> int:
    """
    用 PDF 大纲（书签/TOC）确定哪些 block 是真正的标题。

    PDF 大纲是作者在 Word 里设置的大纲层级（通过 get_toc 获取），
    包含每个标题条目的层级、文本和精确坐标（页码 + y 位置）。
    这是文档结构的权威定义，比 MinerU 的 type 判断可靠得多。

    对每个大纲条目，按坐标匹配 MinerU block，注入 block["_toc_level"]。
    返回成功匹配的大纲条目数。

    下游效果：
      - 有 _toc_level 的 block → 真标题，用大纲层级
      - 无 _toc_level 但 MinerU 标记为 title → 降级为正文（MinerU 误判）
    """
    toc = fitz_doc.get_toc(simple=False)
    if not toc:
        return 0

    # 构建大纲查找表：{(page_idx, y_bucket): level}
    # y_bucket 精度 10pt（容差匹配）
    # 注意：PyMuPDF get_toc 的 to.y 已经是 top-down 坐标（与 MinerU bbox 一致），
    # 不需要做 PDF→top-down 转换
    toc_lookup: dict[tuple[int, int], int] = {}
    for entry in toc:
        level = entry[0]
        dest = entry[3] if len(entry) > 3 else {}
        # entry[2] 是 1-based 页码，dest["page"] 是 0-based 页码
        pi = dest.get("page", entry[2] - 1)
        to = dest.get("to")
        if to is None:
            continue
        y_bucket = int(to.y // 10)
        toc_lookup[(pi, y_bucket)] = level
        # 相邻 bucket 也设上（容差 ±10pt）
        toc_lookup[(pi, y_bucket + 1)] = level

    matched = 0
    for page in pdf_info:
        pi = page.get("page_idx", 0)
        page_blocks = page.get("para_blocks") or page.get("blocks") or []
        for block in page_blocks:
            bbox = block.get("bbox")
            if not (bbox and len(bbox) >= 4):
                continue
            # 用 block 的 y0 匹配大纲
            y_bucket = int(bbox[1] // 10)
            level = toc_lookup.get((pi, y_bucket))
            if level is not None:
                block["_toc_level"] = level
                matched += 1

    return matched


def _extract_table_html(block: dict[str, Any]) -> str:
    """
    从 MinerU table block 提取 HTML。
    MinerU 3.x 的 table HTML 不在 table_body 字段里，
    而是在 blocks[].lines[].spans[].html 里。
    """
    # 路径1：顶层 table_body（部分版本可能有）
    table_body = block.get("table_body") or ""
    if table_body:
        return table_body
    # 路径2：blocks[].lines[].spans[].html（MinerU 3.x 实际位置）
    for sb in block.get("blocks", []):
        for line in sb.get("lines", []):
            for sp in line.get("spans", []):
                h = sp.get("html", "")
                if h:
                    return h
    return ""


def _html_has_merged_cells(table_html: str) -> bool:
    """
    检测 MinerU HTML 是否含合并单元格（colspan>1 或 rowspan>1）。

    用于决定表格重建路径：
    - 含合并单元格 → MinerU HTML 的结构更可靠（AI 视觉能正确识别合并），
      走 HTML 路径。PyMuPDF 的 find_tables() 对含 colspan/rowspan 的复杂表格
      cell 边界检测不可靠（会把每个 cell 都报告成跨满宽，导致全部误判为 colspan=2）。
    - 不含合并单元格 → 走 PyMuPDF 路径（避免 MinerU 的 OCR 错误/续行拆分）。
    """
    if not table_html:
        return False
    for m in re.finditer(r'(?:colspan|rowspan)\s*=\s*["\']?(\d+)', table_html, re.I):
        if int(m.group(1)) > 1:
            return True
    return False


def _apply_table_borders(table) -> None:
    """
    手动给表格添加边框（替代 Table Grid 样式）。

    不用 Table Grid 样式是因为它自带 tblCellMar 默认内边距（左右各5.4pt），
    即使在 tblPr 中设 tblCellMar=0，LibreOffice 仍会用样式的默认值。
    不设 style 时表格无默认内边距，cell 可用宽度 = 列宽全宽。
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = tblBorders.find(qn(f'w:{side}'))
        if b is None:
            b = OxmlElement(f'w:{side}')
            tblBorders.append(b)
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')  # 0.5pt
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')


def _set_table_fixed_width(table, col_widths_pt: list[float]) -> None:
    """
    设置表格为固定列宽布局，直接操作 OOXML。

    python-docx 的 table.autofit=False + cell.width 只设了 tblLayout 和 tcW，
    但 **tblGrid 的 gridCol 仍是创建时的等宽值**，且 tblW 是 auto。
    LibreOffice 在 fixed 布局下优先用 tblGrid 决定列宽，导致 cell.width 被忽略。

    本函数补全三处 OOXML：
      1. tblLayout type=fixed（关闭自动调整）
      2. tblGrid gridCol（列宽定义，渲染器实际依据）
      3. tblW type=dxa（表格总宽度=各列宽之和）
      4. 每个 cell 的 tcW（逐 cell 宽度，与 gridCol 一致）
    """
    from docx.oxml.ns import qn
    from docx.shared import Pt
    from docx.oxml import OxmlElement

    tbl = table._tbl

    # 1. tblLayout type=fixed
    tblPr = tbl.tblPr
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')

    # 2. 更新 tblGrid 的 gridCol（列宽定义）
    # pt → twips(dxa)：1pt = 20 twips
    # 补偿系数 1.014：Windows 安装的 FangSong/TimesNewRoman 字符宽度
    # 比 PDF 内嵌字体平均宽约 1.4%（实测比例 1.0142-1.0143），
    # 不补偿会导致接近 cell 宽度的行在 Word 中多换一行。
    col_widths_dxa = [round(w * 20 * 1.014) for w in col_widths_pt]
    total_dxa = sum(col_widths_dxa)

    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        # 清除旧 gridCol，按新列宽重建
        for gc in tblGrid.findall(qn('w:gridCol')):
            tblGrid.remove(gc)
        for w in col_widths_dxa:
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(w))
            tblGrid.append(gridCol)

    # 3. tblW type=auto（表格宽度由 gridCol 决定，不设固定值）
    #    设固定值(dxa)时，Word 可能因舍入误差认为表格超出内容区而压缩列宽。
    #    auto 模式下 Word 直接用 gridCol 值，列宽精确不被压缩。
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'), 'auto')
    tblW.set(qn('w:w'), '0')

    # 4. 逐 cell 设 tcW（与 gridCol 一致）
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(col_widths_pt):
                cell.width = Pt(col_widths_pt[ci])


def _set_table_row_heights(
    table, row_heights_pt: list[float], rule: str = "exact",
) -> None:
    """
    设置表格每行高度，直接操作 OOXML 的 trHeight。

    row_heights_pt 是从 PDF row bbox 提取的每行高度(pt)。
    rule: 'exact' 或 'atLeast'
      - exact：行高固定为 PDF 值，内容溢出会被裁剪（HTML 路径用）
      - atLeast：行高至少为 PDF 值，内容更多时自动增高（PyMuPDF 重建路径用，
        因为 DOCX 渲染换行点可能与 PDF 不同，exact 会裁剪末尾文字）
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for ri, tr in enumerate(table.rows):
        if ri >= len(row_heights_pt):
            break
        h_pt = row_heights_pt[ri]
        if h_pt <= 0:
            continue
        trPr = tr._tr.get_or_add_trPr()
        # 移除旧的 trHeight
        for old in trPr.findall(qn('w:trHeight')):
            trPr.remove(old)
        trHeight = OxmlElement('w:trHeight')
        # pt → twips(dxa)：1pt = 20 twips
        trHeight.set(qn('w:val'), str(round(h_pt * 20)))
        trHeight.set(qn('w:hRule'), rule)
        trPr.append(trHeight)



def _normalize_text(s: str) -> str:
    """去除空白用于文本匹配（空格/换行差异不影响内容比对）。"""
    return re.sub(r"\s+", "", s or "")


def _enrich_rows_with_pymupdf_text(
    rows: list[list[dict[str, Any]]],
    pymupdf_cells: list[list[str]],
) -> None:
    """
    用 PyMuPDF extract() 的带换行文本，替换 HTML rows 里的扁平文本。

    MinerU HTML 把单元格内多行内容合并成无换行的连续字符串，
    PyMuPDF extract() 保留了原文换行（\\n）。

    两种策略：
      1. 行列完全一致 → 按位置直接替换
      2. 行列不一致（如跨页表格 PyMuPDF 只检测到部分行）→
         按"去空白文本前缀匹配"逐 cell 查找替换，匹配不到的保留 HTML 原文

    匹配规则：HTML cell 去空白文本 == PyMuPDF cell 去空白文本的前 N 字符
    （双向取较短者比较），避免空格/换行差异导致漏匹配。
    """
    n_rows_html = len(rows)
    if not pymupdf_cells:
        return  # 无 PyMuPDF 数据，保留 HTML 原文
    n_rows_pm = len(pymupdf_cells)
    n_cols_pm = len(pymupdf_cells[0]) if pymupdf_cells else 0

    # 策略1：行列完全一致，按位置替换
    if n_rows_pm == n_rows_html:
        # HTML 的列数可能因 colspan 变化，逐行检查
        all_match = True
        for ri, row in enumerate(rows):
            html_cols = sum(c.get("colspan", 1) for c in row)
            if html_cols != n_cols_pm:
                all_match = False
                break
        if all_match:
            for ri, row in enumerate(rows):
                for ci, cell_data in enumerate(row):
                    pm_text = pymupdf_cells[ri][ci] if ci < n_cols_pm else None
                    if pm_text and pm_text.strip():
                        cell_data["text"] = pm_text.strip()
            return

    # 策略2：行列不一致，按文本相似度匹配
    # PyMuPDF 和 MinerU 对同一文字的提取可能有字符差异（如"）"vs"〕"、"繫"vs"系"），
    # 用 difflib.SequenceMatcher 计算整体相似度，容忍少量 OCR 差异。
    import difflib

    pm_texts: list[tuple[str, str]] = []  # (去空白文本, 原始带换行文本)
    for pm_row in pymupdf_cells:
        for pm_text in pm_row:
            if pm_text and pm_text.strip():
                norm = _normalize_text(pm_text)
                if norm:
                    pm_texts.append((norm, pm_text.strip()))

    for row in rows:
        for cell_data in row:
            html_text = cell_data.get("text", "")
            if not html_text:
                continue
            html_norm = _normalize_text(html_text)
            if len(html_norm) < 8:
                continue  # 短文本（如条款号"1.1"）不做模糊匹配，避免误匹配

            best_match = None
            best_ratio = 0.0
            for pm_norm, pm_orig in pm_texts:
                if html_norm == pm_norm:
                    best_match = pm_orig
                    break
                # 整体相似度（容忍 OCR 字符差异）
                ratio = difflib.SequenceMatcher(None, html_norm, pm_norm).ratio()
                # 动态阈值：文本越长容忍度越高
                # ≥30字符用0.8，10-30字符用0.9，避免短文本误匹配
                threshold = 0.8 if len(html_norm) >= 30 else 0.9
                if ratio > threshold and ratio > best_ratio:
                    best_ratio = ratio
                    best_match = pm_orig
            if best_match:
                cell_data["text"] = best_match


def _infer_colspan_rowspan(pymupdf_table) -> list[list[dict[str, Any]]] | None:
    """
    从 PyMuPDF Table 的 cell 物理边界推断 colspan/rowspan。

    PyMuPDF 的 row.cells 返回每个物理 cell 的 bbox 或 None：
    - None → 被上方/左方的 cell 合并占据
    - bbox 跨多行 → rowspan > 1（y 范围覆盖多个行高）
    - bbox 跨多列 → colspan > 1（x 范围覆盖多个列宽）

    返回 rows_data[r][c] = {"text": ..., "rowspan": int, "colspan": int}
    被 None 占据的位置返回 {"text": "", "rowspan": 0, "colspan": 0}（跳过标记）。
    """
    if pymupdf_table is None:
        return None

    rows = list(pymupdf_table.rows)
    n_rows = len(rows)
    if n_rows == 0:
        return None

    # 收集行边界（y0）和列边界（x0）
    row_ys = [r.bbox[1] for r in rows]
    # 列边界遍历**所有行**的 cell x 坐标收集（而非只看第一行）。
    # 原因：第一行可能是跨满宽的合并 cell（colspan=N），只有1个 cell，
    # 从它收集到的列边界只有2个点（左/右），会严重低估列数。
    # 遍历所有行取并集，确保收集到所有列分隔线。
    col_xs = set()
    for row in rows:
        for cell in row.cells:
            if cell:
                col_xs.add(round(cell[0]))
                col_xs.add(round(cell[2]))
    col_xs = sorted(col_xs)
    n_cols = len(col_xs) - 1 if len(col_xs) >= 2 else 1
    # 交叉验证：PyMuPDF 的 col_count 是内部精确计算的最可靠列数
    if pymupdf_table.col_count and pymupdf_table.col_count > n_cols:
        n_cols = pymupdf_table.col_count
    if n_cols < 1:
        return None

    # 构建 rows_data
    rows_data: list[list[dict[str, Any]]] = []
    for ri in range(n_rows):
        row_data: list[dict[str, Any]] = []
        pm_cells = rows[ri].cells
        for ci in range(n_cols):
            if ci < len(pm_cells) and pm_cells[ci] is not None:
                cell_bbox = pm_cells[ci]
                cx0, cy0, cx1, cy1 = cell_bbox
                # 计算 rowspan/colspan：用边界点索引差值。
                # cell 的左/右边界对齐 col_xs 的某个点，上/下边界对齐 row_ys 的某个点，
                # colspan = 右边界索引 - 左边界索引，rowspan = 下边界索引 - 上边界索引。
                # ⚠️ 不能用"覆盖了多少个边界点"（会多算1：cell 两端本身是边界点）。
                def _span_of(val0, val1, ticks, tol=1):
                    """val0/val1 在 ticks 中的索引差 = 跨度"""
                    i = j = None
                    for idx, t in enumerate(ticks):
                        if abs(t - val0) <= tol:
                            i = idx
                        if abs(t - val1) <= tol:
                            j = idx
                    if i is not None and j is not None:
                        return max(1, j - i)
                    return 1
                rowspan = _span_of(cy0, cy1, row_ys)
                colspan = _span_of(cx0, cx1, col_xs)
                row_data.append({"text": "", "rowspan": rowspan, "colspan": colspan})
            else:
                # 被合并的位置
                row_data.append({"text": "", "rowspan": 0, "colspan": 0})
        rows_data.append(row_data)

    # 填入 extract() 的文本（只填非合并位置）
    extracted = pymupdf_table.extract()
    for ri in range(min(n_rows, len(extracted))):
        for ci in range(min(n_cols, len(extracted[ri]))):
            if ri < len(rows_data) and ci < len(rows_data[ri]):
                if rows_data[ri][ci]["rowspan"] > 0:
                    rows_data[ri][ci]["text"] = extracted[ri][ci] or ""

    return rows_data


def _build_table_from_pymupdf(
    doc, block: dict[str, Any], pymupdf_cells: list[list[str]],
) -> None:
    """
    从 PyMuPDF 表格数据直接构建 DOCX 表格。

    pymupdf_cells: PyMuPDF extract() 返回的 cell 文本矩阵 [row][col]
    colspan/rowspan 推断结果从 block["_table_merged_cells"] 读取（预扫描时已计算）。
    """
    n_rows = len(pymupdf_cells)
    n_cols = max((len(row) for row in pymupdf_cells), default=1)
    if n_rows == 0:
        return

    table = doc.add_table(rows=n_rows, cols=n_cols)
    _apply_table_borders(table)

    # 列宽
    col_widths = block.get("_col_widths")
    if col_widths and len(col_widths) == n_cols:
        _set_table_fixed_width(table, col_widths)

    # cell 样式和行高（从 block 中预扫描提取的 PyMuPDF 数据）
    cell_styles = block.get("_table_cell_styles")
    table_lh = block.get("_table_line_height") or block.get("_doc_line_height")
    row_line_heights = block.get("_table_row_line_heights")

    # 构建 rows_data：用预扫描时计算的 colspan/rowspan 推断结果
    # （PyMuPDF Table 对象存在状态共享 bug，不能保存引用后再调用 extract()，
    # 必须在获取 mtable 的瞬间立即计算 _infer_colspan_rowspan 并存储结果）
    inferred = block.get("_table_merged_cells")
    if inferred and len(inferred) == n_rows:
        rows_data = inferred
    else:
        rows_data = []
        for ri in range(n_rows):
            row = []
            for ci in range(min(n_cols, len(pymupdf_cells[ri]) if ri < len(pymupdf_cells) else 0)):
                row.append({"text": pymupdf_cells[ri][ci] or "", "rowspan": 1, "colspan": 1})
            rows_data.append(row)

    # 先设行高，再填 cell
    # PyMuPDF 重建路径用 atLeast 模式：PyMuPDF 的 break 位置基于 PDF 字体宽度，
    # 但 DOCX 渲染时系统字体宽度不同（即使同名），实际行数可能多于 break 数。
    # exact 模式会裁剪多出的行；atLeast 允许行高自适应。
    row_heights = block.get("_table_row_heights")
    # 回归检测补偿：_auto_align_tables 迭代时给表格行高叠加的微调值
    comp = block.get("_row_height_compensation", 0)
    if row_heights and len(row_heights) == n_rows:
        if comp:
            row_heights = [h + comp for h in row_heights]
        _set_table_row_heights(table, row_heights, rule="atLeast")
    elif not row_heights:
        # 行高兜底：用 block bbox 高度 ÷ 行数（同 HTML 路径的兜底逻辑）
        bbox = block.get("bbox") or []
        if len(bbox) >= 4 and n_rows > 0:
            block_h = bbox[3] - bbox[1]
            if block_h > 20:
                avg_h = block_h / n_rows
                _set_table_row_heights(table, [avg_h] * n_rows, rule="atLeast")

    _fill_table_cells(table, rows_data, block.get("_style"), table_lh,
                      row_line_heights, cell_styles)


def _build_table(doc, block: dict[str, Any]) -> None:
    """
    表格重建：根据表格特征选择数据源。

    路径选择策略（由 HTML 合并单元格特征决定）：
    - **含合并单元格（colspan>1 或 rowspan>1）→ MinerU HTML 路径**：
      这类复杂表格 PyMuPDF 的 find_tables() cell 边界检测不可靠（会把每个 cell
      都报告成跨满宽，导致 _infer_colspan_rowspan 全部误判为 colspan=2）。
      MinerU 的 AI 视觉能正确识别合并结构。文本仍用 PyMuPDF 增强（更准确）。
    - **不含合并单元格 → PyMuPDF 路径**（默认）：
      PyMuPDF extract() 是 PDF 结构直接提取，文字 100% 准确、行结构忠实于原 PDF，
      避免 MinerU HTML 的 OCR 错误、续行拆分、文本错位等问题。
    """
    # 跳过已被跨页合并处理的续页 block
    if block.get("_table_merged"):
        return

    pymupdf_cells = block.get("_pymupdf_cells")
    table_html = _extract_table_html(block)

    # 含合并单元格的复杂表格 → HTML 路径（结构更可靠）
    # 但如果 HTML 行数与 PyMuPDF 行数不匹配（MinerU 行结构错误），
    # 放弃 HTML 路径，走 PyMuPDF 路径（_infer_colspan_rowspan 推断合并）
    use_html = False
    if _html_has_merged_cells(table_html):
        html_rows = _parse_html_table(table_html)
        if html_rows and pymupdf_cells:
            # 行数匹配才用 HTML 路径
            if len(html_rows) == len(pymupdf_cells):
                use_html = True
            else:
                print(f"  ⚠️ 表格行数不匹配(HTML={len(html_rows)} vs PyMuPDF={len(pymupdf_cells)})，走PyMuPDF路径",
                      file=sys.stderr)
        elif html_rows and not pymupdf_cells:
            use_html = True  # 无 PyMuPDF 数据，只能用 HTML

    if use_html:
        rows = _parse_html_table(table_html)
        if rows:
            max_cols = max(sum(c.get("colspan", 1) for c in row) for row in rows)
            table = doc.add_table(rows=len(rows), cols=max_cols)
            _apply_table_borders(table)

            col_widths = block.get("_col_widths")
            if col_widths and len(col_widths) == max_cols:
                _set_table_fixed_width(table, col_widths)
            else:
                table.autofit = True

            # 用 PyMuPDF 精确文本增强 HTML 行（PyMuPDF 文字无 OCR 错误）
            _enrich_rows_with_pymupdf_text(rows, pymupdf_cells)

            table_lh = block.get("_table_line_height") or block.get("_doc_line_height")
            row_line_heights = block.get("_table_row_line_heights")
            row_heights = block.get("_table_row_heights")
            if row_heights:
                if len(row_heights) == len(table.rows):
                    _set_table_row_heights(table, row_heights, rule="atLeast")
                elif len(row_heights) > 0:
                    # 行数不匹配（MinerU HTML 行数 ≠ PyMuPDF 行数）：
                    # 按总高度等比分配到每行
                    total_h = sum(row_heights)
                    avg_h = total_h / len(table.rows)
                    _set_table_row_heights(table, [avg_h] * len(table.rows), rule="atLeast")

            _fill_table_cells(table, rows, block.get("_style"), table_lh, row_line_heights,
                              block.get("_table_cell_styles"))
            return

    # 默认：PyMuPDF 重建路径（简单表格，文字更准确）
    if pymupdf_cells:
        _build_table_from_pymupdf(doc, block, pymupdf_cells)
        return

    # 回退：无 PyMuPDF 数据时用 MinerU HTML
    if not table_html:
        return

    rows = _parse_html_table(table_html)
    if not rows:
        return

    max_cols = max(sum(c.get("colspan", 1) for c in row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    _apply_table_borders(table)

    col_widths = block.get("_col_widths")
    if col_widths and len(col_widths) == max_cols:
        _set_table_fixed_width(table, col_widths)
    else:
        table.autofit = True

    _enrich_rows_with_pymupdf_text(rows, pymupdf_cells)
    _fix_crosspage_cell_misplacement(rows, pymupdf_cells)

    table_lh = block.get("_table_line_height") or block.get("_doc_line_height")
    row_line_heights = block.get("_table_row_line_heights")
    row_heights = block.get("_table_row_heights")
    if row_heights and len(row_heights) == len(table.rows):
        _set_table_row_heights(table, row_heights, rule="atLeast")
    elif not row_heights:
        # 行高兜底：PyMuPDF 检测不到的表格（如线条画的空方框/身份证粘贴区）
        # 没有预扫描行高数据，用 block bbox 高度 ÷ 行数 推断。
        # bbox 高度是 MinerU 正确测量的（如83pt方框），可靠。
        bbox = block.get("bbox") or []
        n_rows_tbl = len(table.rows)
        if len(bbox) >= 4 and n_rows_tbl > 0:
            block_h = bbox[3] - bbox[1]
            if block_h > 20:  # 避免极小值
                avg_h = block_h / n_rows_tbl
                _set_table_row_heights(table, [avg_h] * n_rows_tbl, rule="atLeast")

    _fill_table_cells(table, rows, block.get("_style"), table_lh, row_line_heights,
                      block.get("_table_cell_styles"))


def _parse_html_table(table_html: str) -> list[list[dict[str, Any]]]:
    """
    解析 HTML 表格为 [[{text, rowspan, colspan}, ...], ...]。

    ⚠️ MVP 用正则粗解析（够用），生产级建议用 lxml/html.parser。
    TODO(完善): 处理 <thead>/<tbody>、嵌套表、跨页表头重复等边界情况。
    """
    rows: list[list[dict[str, Any]]] = []

    # 按行分割 <tr>...</tr>
    tr_pattern = re.compile(r"<tr[^>]*>(.*?)</tr>", re.DOTALL | re.IGNORECASE)
    # 匹配整个 <td>/<th> 单元格：捕获组1=标签名，组2=属性串，组3=单元格内容
    td_pattern = re.compile(
        r"<(td|th)([^>]*)>(.*?)</\1>",
        re.DOTALL | re.IGNORECASE,
    )
    # 从属性串里提取 rowspan / colspan 值
    rowspan_pat = re.compile(r"rowspan\s*=\s*['\"]?(\d+)", re.IGNORECASE)
    colspan_pat = re.compile(r"colspan\s*=\s*['\"]?(\d+)", re.IGNORECASE)

    for tr_match in tr_pattern.finditer(table_html):
        row: list[dict[str, Any]] = []
        tr_content = tr_match.group(1)
        for td_match in td_pattern.finditer(tr_content):
            attrs = td_match.group(2) or ""
            rowspan_m = rowspan_pat.search(attrs)
            colspan_m = colspan_pat.search(attrs)
            rowspan = int(rowspan_m.group(1)) if rowspan_m else 1
            colspan = int(colspan_m.group(1)) if colspan_m else 1
            # 去除单元格内的 HTML 标签
            cell_html = td_match.group(3)
            cell_text = re.sub(r"<[^>]+>", "", cell_html).strip()
            cell_text = html.unescape(cell_text)  # 反转义所有 HTML 实体
            row.append({"text": cell_text, "rowspan": rowspan, "colspan": colspan})
        if row:
            rows.append(row)

    return rows


def _merge_continuation_rows(
    rows: list[list[dict[str, Any]]],
) -> list[list[dict[str, Any]]]:
    """
    合并 MinerU 跨页表格中的"续行"。

    MinerU 在处理大单元格（内容跨多行）时，有时会把同一行的文字溢出生
    成额外的 <tr> 元素。这些"续行"特征：
      - col0（条款号列）为空
      - col1 文本不以条款号/标题开头（不是新行的开始）

    将它们合并回上一行，避免表格行数膨胀（如原 PDF 22 行 → MinerU 33 行）。
    """
    import re

    # 判断文本是否像一个条款号（如 "1.1"、"5"、"1.18 1.19"）
    def _is_clause_number(text: str) -> bool:
        if not text.strip():
            return False
        # 单个条款号: 数字(可选.数字) + 可选空格
        # 多个条款号空格分隔: 上条重复
        return bool(re.match(r'^\d+(\.\d+)*(\s+\d+(\.\d+)*)*\s*$', text.strip()))

    # 判断是否为表头行（如"条款号"、"内容"）
    def _is_header_like(text: str) -> bool:
        t = text.strip()
        return t in ('条款号', '内容', '名称', '说明', '备注')

    merged: list[list[dict[str, Any]]] = []
    for ri, row in enumerate(rows):
        if not row:
            merged.append(row)
            continue

        is_continuation = False
        if merged and len(row) >= 2:
            col0_text = (row[0].get("text", "") or "").strip()
            # 续行：col0 不像条款号（非 "1.1"、"5"、"1.18 1.19" 格式）
            # 且排除表头行（如"条款号"）——它们虽然不像条款号，但是合法的第一列文本
            if not _is_clause_number(col0_text) and not _is_header_like(col0_text):
                is_continuation = True

        if is_continuation:
            prev_row = merged[-1]
            for ci, cell in enumerate(row):
                if ci < len(prev_row):
                    cur_text = cell.get("text", "") or ""
                    if cur_text.strip():
                        # col0 是条款号列，续行的 col0 通常是溢出文本（如"商自行承担"），
                        # 不应合并到上一行的条款号中。只合并内容列（ci >= 1）的文本。
                        if ci == 0:
                            continue
                        prev_text = prev_row[ci].get("text", "") or ""
                        prev_row[ci]["text"] = prev_text + cur_text
        else:
            merged.append([dict(c) for c in row])

    return merged


def _fix_crosspage_cell_misplacement(
    rows: list[list[dict[str, Any]]],
    pymupdf_cells: list[list[str]] | None = None,
) -> None:
    """
    修复 MinerU 跨页表格的单元格错位。

    问题：跨页表格在换页处，MinerU 常把内容列的文字错放到条款号列（col0）。
    例如条款号列出现 "商自行承担。应商自行承担。可能导致投标..."（整句正文）。

    修复逻辑（两层）：
      1. PyMuPDF 数据覆盖的行：对比 col0，PyMuPDF 显示为空但 HTML 有文本 → 错位
      2. PyMuPDF 数据未覆盖的行（行数不匹配）：判断 col0 文本长度
         —— 条款号通常 ≤10 字符，col0 超过 20 字符几乎一定是正文错位
    """
    if not pymupdf_cells:
        return  # 无 PyMuPDF 数据时不做修正

    for ri, row in enumerate(rows):
        if len(row) < 2:
            continue
        col0_text = row[0].get("text", "").strip()
        if not col0_text:
            continue

        is_misplaced = False

        if ri < len(pymupdf_cells):
            # 策略 1：PyMuPDF 数据覆盖到的行，对比 col0
            pm_col0 = ""
            if len(pymupdf_cells[ri]) > 0:
                pm_col0 = (pymupdf_cells[ri][0] or "").strip()
            if not pm_col0 and col0_text:
                is_misplaced = True
        else:
            # 策略 2：PyMuPDF 未覆盖的行，用文本特征判断
            # 条款号特征：短（≤10字符）、无中文标点、通常是数字+点/顿号格式
            # 正文特征：长（>20字符）或含中文标点（句号、逗号等）
            import re as _re
            has_cn_punct = bool(_re.search(r'[。，；！？、：]', col0_text))
            if len(col0_text) > 20 or has_cn_punct:
                is_misplaced = True

        if is_misplaced:
            col1_text = row[1].get("text", "").strip()
            if col1_text:
                row[1]["text"] = col0_text + "\n" + col1_text
            else:
                row[1]["text"] = col0_text
            row[0]["text"] = ""


def _fill_table_cells(table, rows: list[list[dict[str, Any]]],
                      style: dict[str, Any] | None = None,
                      line_height: float | None = None,
                      row_line_heights: list[float | None] | None = None,
                      cell_styles: list[list[list[dict[str, Any]]]] | None = None,
                      natural_wrap: bool = False) -> None:
    """
    填充表格单元格，处理 rowspan/colspan 合并，并应用样式。

    style: block._style（兜底样式，当 cell_styles 不可用时使用）
    line_height: 表格级兜底行高（pt）
    row_line_heights: 每行的文字行间距（pt），优先级最高
    cell_styles: 从 PDF dict 提取的每 cell 内 span 样式，优先级最高。
      格式: cell_styles[r][c] = [{"text","font","size","bold","italic","color"}, ...]
    natural_wrap: True 时不按 PDF 视觉行硬换行（add_break），让 Word 根据
      实际列宽自然换行。用于 PyMuPDF 重建路径——DOCX 字体度量与 PDF 内嵌
      字体不完全一致，硬换行会导致一行放不下时多换行+末尾字被裁剪。
    """
    occupied: dict[tuple[int, int], bool] = {}
    from docx.shared import Pt as _Pt
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml.ns import qn as _qn

    # 获取每行的 trHeight（用于约束 line_spacing，防止内容溢出）
    # 注意：必须与 table.rows 一一对应（含 None），不能跳过无 trHeight 的行，
    # 否则索引错位会导致 safe_lh 计算错误（如合并行读到下一行的 trHeight）。
    row_tr_heights: list[float | None] = []
    for tr in table.rows:
        tr_h = None
        trPr = tr._tr.find(_qn('w:trPr'))
        if trPr is not None:
            th = trPr.find(_qn('w:trHeight'))
            if th is not None:
                tr_h = int(th.get(_qn('w:val'))) / 20
        row_tr_heights.append(tr_h)

    for r_idx, row in enumerate(rows):
        c_idx = 0
        r_lh = line_height or 22
        if row_line_heights and r_idx < len(row_line_heights) and row_line_heights[r_idx]:
            r_lh = row_line_heights[r_idx]

        # 安全约束：根据本行 trHeight 和 cell 文字行数，确保内容不溢出
        # line_spacing 不能超过 trHeight / (max_cell_lines + 0.3)
        # +0.3 给 cell 上下内边距留余量，与 _extract_table_row_geometry 的计算一致
        tr_h = row_tr_heights[r_idx] if r_idx < len(row_tr_heights) else None
        if tr_h and tr_h > 0:
            # 扫描本行所有 cell 的文字行数，取最大值
            max_lines = 0
            for cd in row:
                t_lines = (cd.get("text") or "").count("\n") + 1
                max_lines = max(max_lines, t_lines)
            if max_lines > 0:
                safe_lh = tr_h / (max_lines + 0.5)
                if r_lh > safe_lh:
                    r_lh = safe_lh

        for cell_data in row:
            while occupied.get((r_idx, c_idx)):
                c_idx += 1
            if c_idx >= len(table.columns):
                break

            rowspan = cell_data.get("rowspan", 1)
            colspan = cell_data.get("colspan", 1)
            text = cell_data.get("text", "")

            # rowspan=0/colspan=0 是被合并占位的格子，跳过（不做任何操作）
            # 这些位置的 occupied 标记已由前驱 cell 的 colspan/rowspan 设置
            if rowspan == 0 or colspan == 0:
                c_idx += 1
                continue

            try:
                target_cell = table.cell(r_idx, c_idx)
                target_cell.text = ""  # 清空
                # 设置 cell 内边距：
                # 上下清零（消除 exact 模式下内容被裁剪）；
                # 左右设 28 twips(1.4pt) 而非 0——全0时 WPS 会把文字紧贴
                # cell 边缘导致末尾字溢出，留微小内边距更安全。
                from docx.oxml import OxmlElement as _OxmlEl
                tcPr = target_cell._tc.get_or_add_tcPr()
                tcMar = _OxmlEl('w:tcMar')
                for side in ('top', 'bottom'):
                    m = _OxmlEl(f'w:{side}')
                    m.set(_qn('w:w'), '0')
                    m.set(_qn('w:type'), 'dxa')
                    tcMar.append(m)
                for side in ('left', 'right'):
                    m = _OxmlEl(f'w:{side}')
                    m.set(_qn('w:w'), '28')  # 1.4pt
                    m.set(_qn('w:type'), 'dxa')
                    tcMar.append(m)
                tcPr.append(tcMar)
                para = target_cell.paragraphs[0]
                pf = para.paragraph_format
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = _Pt(r_lh)
                pf.space_before = _Pt(0)
                pf.space_after = _Pt(0)

                # 优先用 PDF per-cell span 序列按实际字体/字号逐个渲染
                spans_data = None
                if (cell_styles and r_idx < len(cell_styles) and
                        c_idx < len(cell_styles[r_idx])):
                    spans_data = cell_styles[r_idx][c_idx]

                if spans_data:
                    # 按 span 的实际样式逐个创建 run，还原 PDF 原始字体/字号
                    # 换行检测：相邻 span y 差 >= 5pt 视为新视觉行
                    # （与 _extract_table_cell_styles 的分组逻辑一致）
                    prev_y = None
                    for sp in spans_data:
                        sp_text = sp.get("text") or ""
                        if not sp_text.strip():
                            continue
                        # 去掉 span 文本开头/末尾的空格
                        # PyMuPDF 在不同字体 span 间会插入空格作为分隔符，
                        # 这个空格在 CJK 字体下占全角宽度（12pt），
                        # 导致行总宽度超出 cell 而触发换行
                        sp_text = sp_text.strip()
                        sp_y = sp.get("_y")
                        if prev_y is not None and sp_y is not None:
                            if abs(sp_y - prev_y) >= 5 and not natural_wrap:
                                # 新视觉行：添加换行符（natural_wrap 模式下跳过，
                                # 让 Word 根据实际列宽自然换行）
                                br_run = para.add_run()
                                br_run.add_break()
                        prev_y = sp_y
                        sp_style = {
                            "font": sp.get("font"),
                            "size": sp.get("size"),
                            "bold": sp.get("bold"),
                            "italic": sp.get("italic"),
                            "color": sp.get("color"),
                            "underline": sp.get("underline", False),
                        }
                        run = para.add_run(sp_text)
                        _apply_style(run, sp_style)
                else:
                    # 无 per-cell 样式数据，回退到 block 级样式
                    if "\n" in text:
                        lines = text.split("\n")
                        run = para.add_run()
                        for li, line in enumerate(lines):
                            if li > 0:
                                run.add_break()
                            if line:
                                run.add_text(line)
                    else:
                        run = para.add_run(text)
                    if style:
                        _apply_style(run, style)

                # 处理 colspan（水平合并）
                if colspan > 1:
                    end_col = min(c_idx + colspan - 1, len(table.columns) - 1)
                    if end_col > c_idx:
                        merge_cell = table.cell(r_idx, end_col)
                        target_cell = target_cell.merge(merge_cell)

                # 处理 rowspan（垂直合并）
                if rowspan > 1:
                    end_row = min(r_idx + rowspan - 1, len(table.rows) - 1)
                    if end_row > r_idx:
                        merge_cell = table.cell(end_row, c_idx)
                        target_cell = target_cell.merge(merge_cell)

                # 标记被占用的位置
                for dr in range(rowspan):
                    for dc in range(colspan):
                        occupied[(r_idx + dr, c_idx + dc)] = True

            except (IndexError, Exception) as e:
                # 超出表格范围的单元格，跳过
                # TODO(完善): 真实 PDF 上观察是否需要动态扩展表格
                print(f"  ⚠️ 表格单元格填充跳过 ({r_idx},{c_idx}): {e}",
                      file=sys.stderr)
                pass

            c_idx += colspan


def _build_image(doc, block: dict[str, Any], images_dir: str | None) -> None:
    """
    图片：从 images/ 读图插入。
    图片宽度从 block 的 bbox 提取（pt→inch），保持与 PDF 原始尺寸一致。
    超出页面内容宽度时等比缩放。
    """
    from docx.shared import Inches

    img_path = block.get("img_path") or block.get("image_path") or ""
    if not img_path or not images_dir:
        return

    full_path = Path(images_dir) / Path(img_path).name
    if not full_path.exists():
        print(f"  ⚠️ 图片不存在: {full_path}", file=sys.stderr)
        return

    # 从 bbox 提取图片宽度（pt），转换为英寸（1 inch = 72pt）
    bbox = block.get("bbox", [])
    width_in = None
    if bbox and len(bbox) >= 4:
        width_pt = bbox[2] - bbox[0]
        if width_pt > 10:  # 过滤异常值
            width_in = width_pt / 72.0

    # 上限保护：不超过页面内容宽度
    page_width = block.get("_page_width", 595)
    page_x0 = block.get("_page_x0", 72)
    content_width_in = (page_width - page_x0 * 2) / 72.0  # 粗估内容宽度
    if width_in is None:
        width_in = content_width_in  # 无 bbox 时用内容宽度
    elif width_in > content_width_in:
        width_in = content_width_in  # 超出时截断到内容宽度

    doc.add_picture(str(full_path), width=Inches(width_in))


def _build_list(doc, block: dict[str, Any]) -> None:
    """
    列表：MVP 降级为普通段落（带编号文本）。
    TODO(完善): 根据 MinerU 的 list sub_type 还原编号/项目符号格式。
    """
    line_height = block.get("_line_height") or block.get("_doc_line_height")
    for line in block.get("lines", []):
        text = "".join(
            s.get("content") or s.get("text") or ""
            for s in line.get("spans", [])
        )
        if text.strip():
            try:
                para = doc.add_paragraph(text, style="List Bullet")
            except KeyError:
                para = doc.add_paragraph(text)
            # 行距：从 PDF 测量的真实行高
            if line_height:
                from docx.shared import Pt
                from docx.enum.text import WD_LINE_SPACING
                pf = para.paragraph_format
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = Pt(line_height)


def _build_index(doc, block: dict[str, Any]) -> None:
    """
    目录（MinerU 的 type=index）：每行一个章节条目，独立段落输出。
    保留行结构，避免所有条目被拼成一个段落。
    TODO: TOC 跳转（书签/超链接）在 MinerU 的 index block 里不保留，无法还原。
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for line in block.get("lines", []):
        text = "".join(
            s.get("content") or s.get("text") or ""
            for s in line.get("spans", [])
        )
        if not text.strip():
            continue
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Pt(0)
        pf.line_spacing = Pt(block.get("_line_height") or block.get("_doc_line_height", 22))
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        # 逐 span 建 run，保留样式
        for sp in line.get("spans", []):
            content = sp.get("content") or sp.get("text") or ""
            if not content:
                continue
            run = para.add_run(content)
            _apply_style(run, sp.get("_style"))


# ═══════════════════════════════════════════════════════════════
#  对外主接口
# ═══════════════════════════════════════════════════════════════

# block type → 重建函数的分发表
# 兼容 MinerU 2.0 多种可能的类型名
_BLOCK_BUILDERS = {
    "title": _build_title,
    "text": _build_text,
    "paragraph": _build_text,      # 别名
    "table": _build_table,
    "image": _build_image,
    "picture": _build_image,       # 别名
    "list": _build_list,
    "index": _build_index,         # 目录（每行独立段落）
}


def _detect_page_margins(pdf_path: str, pdf_info: list) -> dict[str, float] | None:
    """
    用 PyMuPDF 精确提取 PDF 页面边距。
    取多页（跳过首页封面）所有 span 的坐标范围，算出上下左右边距。
    比 MinerU block bbox 精确（MinerU bbox 底部系统性偏低）。

    页脚/页眉排除：页脚（如"—1—"）位于页面最底部，如果不排除，
    max(y1s) 会取到页脚坐标，导致下边距被严重低估（如 47pt vs 真实 84pt）。
    用 y 阈值（页高 ±60pt 范围内的行视为页眉/页脚）排除。
    """
    try:
        import fitz
    except ImportError:
        return None

    try:
        doc = fitz.open(pdf_path)
    except Exception:
        return None

    page_count = doc.page_count
    top_margins = []
    bottom_margins = []
    left_margins = []
    right_margins = []

    # 取第2-10页采样（跳过首页封面）
    for pi in range(1, min(10, page_count)):
        page = doc[pi]
        d = page.get_text("dict")
        pw = page.rect.width
        ph = page.rect.height
        y0s, y1s, x0s, x1s = [], [], [], []
        for block in d.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                ly0 = line["bbox"][1]
                # 排除页脚（页面底部 60pt 范围内）和页眉（顶部 60pt 范围）
                if ly0 > ph - 60 or ly0 < 60:
                    continue
                for sp in line.get("spans", []):
                    if not sp["text"].strip():
                        continue
                    bx = sp["bbox"]
                    y0s.append(bx[1])
                    y1s.append(bx[3])
                    x0s.append(bx[0])
                    x1s.append(bx[2])
        if y0s and y1s:
            # 上下左右边距：都要同时考虑文字 span 和表格 bbox 的边界。
            top_min = min(y0s)
            bottom_max = max(y1s)
            left_min = min(x0s)
            right_max = max(x1s)
            try:
                for tab in page.find_tables().tables:
                    if tab.bbox[1] >= ph - 60:
                        continue
                    top_min = min(top_min, tab.bbox[1])
                    bottom_max = max(bottom_max, tab.bbox[3])
                    left_min = min(left_min, tab.bbox[0])
                    right_max = max(right_max, tab.bbox[2])
            except Exception:
                pass
            top_margins.append(top_min)
            bottom_margins.append(ph - bottom_max)
            left_margins.append(left_min)
            right_margins.append(pw - right_max)

    doc.close()

    if not top_margins:
        return None

    top_margins.sort()
    bottom_margins.sort()
    left_margins.sort()
    right_margins.sort()

    return {
        # 上下左右边距都取最小值：内容最靠四边的页反映真实可用内容区。
        "top": min(top_margins),
        "bottom": min(bottom_margins),
        "left": min(left_margins),
        "right": min(right_margins),
    }


def _detect_per_page_margins(
    pdf_path: str, pdf_info: list,
) -> dict[int, dict[str, float]] | None:
    """
    为每个 page_idx 独立计算 PDF 边距（page_idx → {top, bottom, left, right}）。

    每页的边距由该页 span + 表格 bbox 的极值决定：
    - 上下边距：合并表格 bbox（表格高度占满页面时保证内容区高度）
    - 左右边距：只用 span，不合并表格 bbox（表格有自己的 _col_widths，
      不靠页面边距撑宽度；合并会把正文行拉宽）
    """
    try:
        import fitz
    except ImportError:
        return None

    try:
        doc = fitz.open(pdf_path)
    except Exception:
        return None

    result: dict[int, dict[str, float]] = {}
    # 全局上下边距：取所有页最小值（避免封面页的大留白被当成边距）
    global_top = []
    global_bottom = []

    for page in pdf_info:
        pi = page.get("page_idx", 0)
        if pi >= doc.page_count:
            continue
        fpage = doc[pi]
        d = fpage.get_text("dict")
        pw = fpage.rect.width
        ph = fpage.rect.height
        y0s, y1s, x0s, x1s = [], [], [], []
        for block in d.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                ly0 = line["bbox"][1]
                if ly0 > ph - 60 or ly0 < 60:
                    continue
                for sp in line.get("spans", []):
                    if not sp["text"].strip():
                        continue
                    bx = sp["bbox"]
                    y0s.append(bx[1])
                    y1s.append(bx[3])
                    x0s.append(bx[0])
                    x1s.append(bx[2])
        if not y0s:
            continue  # 空页跳过

        # 判断是否纯表格页：MinerU 数据中该页只有 table block
        mu_page = next((p for p in pdf_info if isinstance(p, dict) and p.get("page_idx") == pi), {})
        mu_blocks = mu_page.get("para_blocks") or mu_page.get("blocks") or []
        is_table_only = bool(mu_blocks) and all(
            (b.get("type") or "").lower() == "table" for b in mu_blocks
        )

        top_min = min(y0s)
        bottom_max = max(y1s)
        left_min = min(x0s)
        right_max = max(x1s)
        try:
            tab_bboxes = [tab.bbox for tab in fpage.find_tables().tables
                          if tab.bbox[1] < ph - 60]
            for tb in tab_bboxes:
                top_min = min(top_min, tb[1])
                bottom_max = max(bottom_max, tb[3])
            # 纯表格页（MinerU 只有 table block）：左右边距用表格 bbox
            # 否则少数正文 span 位置会给出偏大的边距（如标题居中 left=188）
            if is_table_only and tab_bboxes:
                for tb in tab_bboxes:
                    left_min = min(left_min, tb[0])
                    right_max = max(right_max, tb[2])
        except Exception:
            pass

        global_top.append(top_min)
        global_bottom.append(ph - bottom_max)
        result[pi] = {
            "left": round(left_min, 1),
            "right": round(pw - right_max, 1),
        }

    doc.close()

    # 上下边距用全局最小值（内容最靠边的页反映真实可用内容区）
    gt = min(global_top) if global_top else 72
    gb = min(global_bottom) if global_bottom else 72
    for pi in result:
        result[pi]["top"] = round(gt, 1)
        result[pi]["bottom"] = round(gb, 1)

    return result if result else None


def build_docx(
    merged_data: dict[str, Any],
    images_dir: str | None,
    output_path: str | Path,
    pdf_path: str | None = None,
) -> str:
    """
    遍历合并后的 blocks，按 type 分发重建 DOCX。

    参数：
        merged_data:  align.align_and_merge 的输出（含 _style 的 MinerU 数据）
        images_dir:   图片目录路径（可能为 None）
        output_path:  输出 docx 路径
        pdf_path:     源 PDF 路径（用于精确提取页面边距，可选）

    返回：
        输出文件的绝对路径。

    参数：
        merged_data:  align.align_and_merge 的输出（含 _style 的 MinerU 数据）
        images_dir:   图片目录路径（可能为 None）
        output_path:  输出 docx 路径

    返回：
        输出文件的绝对路径。

    处理顺序：按 page_idx 顺序遍历，每页内按 block 顺序。
    利用 block 的 bbox Y 坐标差还原垂直间距（留白）。
    """
    from docx import Document
    from docx.shared import Mm, Pt

    doc = Document()
    pdf_info = merged_data.get("pdf_info", [])

    # ── 页面设置：从 PDF 真实尺寸设置纸张大小和边距 ──
    # python-docx 默认用 US Letter，需改为 PDF 的实际尺寸（通常是 A4）
    # 每页独立边距：不同页面（正文页 vs 表格页）有不同的内容区域宽度
    per_page_margins = _detect_per_page_margins(pdf_path, pdf_info) if pdf_path else None

    if pdf_info:
        first_page = pdf_info[0]
        page_size = first_page.get("page_size", [595, 842])
        pdf_w_pt = page_size[0] if len(page_size) > 0 else 595
        pdf_h_pt = page_size[1] if len(page_size) > 1 else 842

        sec = doc.sections[0]
        # 方向判断（支持横版首页）
        from docx.enum.section import WD_ORIENT
        if pdf_w_pt > pdf_h_pt:
            sec.orientation = WD_ORIENT.LANDSCAPE
        else:
            sec.orientation = WD_ORIENT.PORTRAIT
        sec.page_width = Mm(pdf_w_pt * 0.3528)
        sec.page_height = Mm(pdf_h_pt * 0.3528)

        # 首页边距优先用 per_page_margins，回退到全局 margins
        if per_page_margins and 0 in per_page_margins:
            pm = per_page_margins[0]
            sec.left_margin = Pt(pm["left"])
            sec.right_margin = Pt(pm["right"])
            sec.top_margin = Pt(pm["top"])
            sec.bottom_margin = Pt(pm["bottom"])
        elif per_page_margins:
            # 首页无数据，用第一个有效页的边距
            pm = next(iter(per_page_margins.values()))
            sec.left_margin = Pt(pm["left"])
            sec.right_margin = Pt(pm["right"])
            sec.top_margin = Pt(pm["top"])
            sec.bottom_margin = Pt(pm["bottom"])
        else:
            # 无 pdf_path 时从 MinerU block bbox 推断
            first_blocks = first_page.get("para_blocks") or first_page.get("blocks") or []
            if first_blocks:
                all_x0 = [b["bbox"][0] for b in first_blocks if b.get("bbox")]
                all_x1 = [b["bbox"][2] for b in first_blocks if b.get("bbox")]
                if all_x0 and all_x1:
                    sec.left_margin = Pt(min(all_x0))
                    sec.right_margin = Pt(pdf_w_pt - max(all_x1))
            if not sec.top_margin:
                sec.top_margin = Pt(72)
            if not sec.bottom_margin:
                sec.bottom_margin = Pt(72)

    # ── 跨页 block 拆分：按 PDF 页边界切开，复刻 PDF 分页 ──
    # MinerU 把跨页段落塞进一个 block，bbox 只覆盖本页。
    # 拆成本页部分 + 续页部分（替换下一页空占位），配合分页符实现分页复刻。
    # ⚠️ 必须在行高提取之前执行：拆分后各部分 lines 数正确，bbox高/行数 才能算准。
    split_count = _split_crosspage_blocks(pdf_info)
    if split_count:
        print(f"  📄 拆分 {split_count} 个跨页 block（本页+续页分离，插分页符）",
              file=sys.stderr)

    # ── 预扫描：从 PDF 提取排版信息（行高、列宽、单元格换行）──
    # 对所有 block 类型提取真实行高（block["_line_height"]），
    # 使后续不同 PDF 都能各自还原真实行距，而非用统一默认值。
    # table 类型额外提取列宽和带换行的单元格文本。
    if pdf_path:
        try:
            import fitz
            fdoc = fitz.open(pdf_path)
            line_h_found = 0
            widths_found = 0
            texts_found = 0
            for page in pdf_info:
                pi = page.get("page_idx", 0)
                if pi >= fdoc.page_count:
                    continue
                page_blocks = page.get("para_blocks") or page.get("blocks") or []
                for block in page_blocks:
                    btype = (block.get("type") or block.get("block_type") or "").lower()
                    bbox = block.get("bbox")
                    if not (bbox and len(bbox) >= 4):
                        continue

                    # 所有类型：提取真实行高（bbox高度 / MinerU行数）
                    # table 类型跳过——table block 的 bbox 是整个表格的高度，
                    # MinerU 行数(lines)为 0，bbox高/1 = 表格总高（如 600pt），
                    # 传给单元格行距会导致每行占满一页。
                    if btype != "table":
                        n_lines = len(block.get("lines", [])) or 1
                        line_h = _measure_block_line_height(bbox, n_lines)
                        if line_h:
                            block["_line_height"] = line_h
                            line_h_found += 1

                    # table 类型：额外提取列宽、单元格文本、行高
                    if btype == "table":
                        fitz_page = fdoc[pi]
                        mtable = _find_matching_pymupdf_table(fdoc, pi, bbox)
                        if mtable is None:
                            continue
                        widths = _extract_table_col_widths_from_table(mtable)
                        if widths:
                            block["_col_widths"] = widths
                            widths_found += 1
                        cell_texts = _extract_table_cell_texts(mtable)
                        if cell_texts:
                            block["_pymupdf_cells"] = cell_texts
                            # 立即推断 colspan/rowspan（PyMuPDF Table 对象存在状态共享 bug，
                            # 对其他页的 Table 调用任何操作会污染本页 Table 的 extract() 结果，
                            # 所以必须在获取 mtable 的瞬间立即计算，不存 Table 引用）
                            block["_table_merged_cells"] = _infer_colspan_rowspan(mtable)
                            texts_found += 1
                        cell_styles = _extract_table_cell_styles(mtable, fitz_page)
                        if cell_styles:
                            block["_table_cell_styles"] = cell_styles
                        # 行高和 cell 内文字行高（从 PDF 原始数据提取）
                        row_heights, text_lh, row_lhs = _extract_table_row_geometry(mtable, fitz_page)
                        if row_heights:
                            # 防溢出：表格总高超过该页可用高度时等比压缩。
                            # 用该页的实际边距（从 per_page_margins 获取），而非全局默认值。
                            total_h = sum(row_heights)
                            page_h = fitz_page.rect.height
                            pm = per_page_margins.get(pi, {}) if per_page_margins else {}
                            top_m = pm.get("top", 72)
                            bot_m = pm.get("bottom", 54)
                            avail_h = page_h - top_m - bot_m - 30  # 5pt 渲染余量
                            if total_h > avail_h:
                                scale = avail_h / total_h
                                row_heights = [round(h * scale, 1) for h in row_heights]
                                row_lhs = [round(lh * scale, 1) if lh else None
                                           for lh in row_lhs]
                            block["_table_row_heights"] = row_heights
                            block["_table_row_line_heights"] = row_lhs
                        if text_lh:
                            block["_table_line_height"] = text_lh

            # ── 跨页表格数据分发（替代旧的合并逻辑）──
            # MinerU 把跨页表格的全部行放在主 block（page_idx=N）的 HTML 中，
            # 续页 block（page_idx=N+1...）是空占位（lines_deleted, 无HTML）。
            # 旧逻辑把续页数据合并到主block → 33行连续表格 → DOCX自行分页（不可控偏移）。
            # 新逻辑：主block只保留本页行数，每个续页block独立填充 → 每页独立表格+分页符。
            for pi, page in enumerate(pdf_info):
                page_idx = page.get("page_idx", 0)
                if page_idx >= fdoc.page_count:
                    continue
                page_blocks = page.get("para_blocks") or page.get("blocks") or []
                for bi, block in enumerate(page_blocks):
                    btype = (block.get("type") or "").lower()
                    if btype != "table":
                        continue
                    html = _extract_table_html(block)
                    if not html:
                        continue  # 续页 block 无 HTML，跳过（在下方被主block分发处理）
                    # 此 block 有 HTML → 跨页表格的主 block
                    # 本页 PyMuPDF 行数（主block本页应保留的行数）
                    main_cells = block.get("_pymupdf_cells") or []
                    main_heights = block.get("_table_row_heights") or []
                    main_row_lhs = block.get("_table_row_line_heights") or []
                    main_styles = block.get("_table_cell_styles") or []
                    n_main = len(main_cells)  # 本页行数

                    # 主block截断：只保留本页数据（已在预扫描中提取，就是 main_cells）
                    # main_cells 已经是本页的行数据（预扫描只提取了本页），不需要截断

                    # 向后扫描续页，给每个续页block填充独立数据
                    for look_pi in range(pi + 1, len(pdf_info)):
                        look_page = pdf_info[look_pi]
                        look_idx = look_page.get("page_idx", 0)
                        if look_idx >= fdoc.page_count:
                            continue
                        found_continuation = False
                        for look_block in (look_page.get("para_blocks") or look_page.get("blocks") or []):
                            if (look_block.get("type") or "").lower() != "table":
                                continue
                            look_html = _extract_table_html(look_block)
                            if look_html:
                                break  # 另一个有HTML的table → 不是续页
                            # 续页 block：填充本页独立数据（不合并到主block）
                            look_bbox = look_block.get("bbox")
                            if not (look_bbox and len(look_bbox) >= 4):
                                continue
                            mtable = _find_matching_pymupdf_table(fdoc, look_idx, look_bbox)
                            if mtable is None:
                                continue
                            ct = _extract_table_cell_texts(mtable)
                            if ct:
                                look_block["_pymupdf_cells"] = ct
                                # 立即推断合并单元格（PyMuPDF状态共享bug，必须即时计算）
                                look_block["_table_merged_cells"] = _infer_colspan_rowspan(mtable)
                            fitz_page = fdoc[look_idx]
                            cs = _extract_table_cell_styles(mtable, fitz_page)
                            if cs:
                                look_block["_table_cell_styles"] = cs
                            rh, tlh, row_lhs = _extract_table_row_geometry(mtable, fitz_page)
                            if rh:
                                # 防溢出：同主 block 逻辑，用该页实际边距
                                total_h = sum(rh)
                                page_h = fitz_page.rect.height
                                pm = per_page_margins.get(look_idx, {}) if per_page_margins else {}
                                avail_h = page_h - pm.get("top", 72) - pm.get("bottom", 54) - 30
                                if total_h > avail_h:
                                    scale = avail_h / total_h
                                    rh = [round(h * scale, 1) for h in rh]
                                    row_lhs = [round(lh * scale, 1) if lh else None
                                               for lh in row_lhs]
                                look_block["_table_row_heights"] = rh
                                look_block["_table_row_line_heights"] = row_lhs
                            if tlh:
                                look_block["_table_line_height"] = tlh
                            # 列宽
                            widths = _extract_table_col_widths_from_table(mtable)
                            if widths:
                                look_block["_col_widths"] = widths
                            # 不标记 _table_merged（让它独立渲染）
                            # 标记为续页表格（触发分页符）
                            look_block["_crosspage_continuation"] = True
                            found_continuation = True
                        if not found_continuation:
                            break  # 没找到续页 table → 停止扫描

                    # 主block也标记为跨页表格（用于分页逻辑判断）
                    block["_table_crosspage"] = True

            # ── 补救 MinerU 漏检的表格 ──
            # MinerU 有时把整个表格误判为多个 text block（如第13页），
            # 导致 DOCX 中丢失表格结构。用 PyMuPDF find_tables 交叉校验：
            # 如果 PyMuPDF 在某页检测到表格但 MinerU 没有 table block，
            # 将该页 text block 替换为一个 table block（含 PyMuPDF 数据）。
            for page in pdf_info:
                pi = page.get("page_idx", 0)
                if pi >= fdoc.page_count:
                    continue
                fitz_page = fdoc[pi]
                # MinerU 是否已有 table block
                has_table = any(
                    (b.get("type") or "").lower() == "table"
                    for b in (page.get("para_blocks") or page.get("blocks") or [])
                )
                if has_table:
                    continue  # 已有表格，不需要补救

                # PyMuPDF 是否检测到表格
                pm_tabs = fitz_page.find_tables()
                for pm_t in pm_tabs.tables:
                    # 验证是真表格（有边框/足够行数），避免误判
                    if pm_t.row_count < 1 or pm_t.col_count < 2:
                        continue
                    cell_texts = _extract_table_cell_texts(pm_t)
                    if not cell_texts or not any(
                        any(c and c.strip() for c in row) for row in cell_texts
                    ):
                        continue  # 所有 cell 为空，不是真表格
                    widths = _extract_table_col_widths_from_table(pm_t)
                    row_heights, text_lh, row_lhs = _extract_table_row_geometry(pm_t, fitz_page)
                    cell_styles = _extract_table_cell_styles(pm_t, fitz_page)

                    # 创建一个新的 table block
                    new_block: dict[str, Any] = {
                        "type": "table",
                        "bbox": list(pm_t.bbox),
                        "blocks": [],
                        "lines": [],
                        "_style": {"font": "FangSong", "size": 12.0, "bold": False, "italic": False, "color": 0},
                    }
                    if cell_texts:
                        new_block["_pymupdf_cells"] = cell_texts
                        # 立即推断 colspan/rowspan（同预扫描逻辑，不存 Table 引用）
                        new_block["_table_merged_cells"] = _infer_colspan_rowspan(pm_t)
                    if widths:
                        new_block["_col_widths"] = widths
                    if row_heights:
                        # 防溢出：同主 block 逻辑，用该页实际边距
                        total_h = sum(row_heights)
                        page_h = fitz_page.rect.height
                        ppm = per_page_margins.get(pi, {}) if per_page_margins else {}
                        avail_h = page_h - ppm.get("top", 72) - ppm.get("bottom", 54) - 30
                        if total_h > avail_h:
                            scale = avail_h / total_h
                            row_heights = [round(h * scale, 1) for h in row_heights]
                            row_lhs = [round(lh * scale, 1) if lh else None
                                       for lh in row_lhs]
                        new_block["_table_row_heights"] = row_heights
                        new_block["_table_row_line_heights"] = row_lhs
                    if text_lh:
                        new_block["_table_line_height"] = text_lh
                    if cell_styles:
                        new_block["_table_cell_styles"] = cell_styles

                    # 替换该页 blocks：移除落在表格 bbox 内的 text block
                    tx0, ty0, tx1, ty1 = pm_t.bbox
                    old_blocks = page.get("para_blocks") or page.get("blocks") or []
                    kept = []
                    removed = 0
                    for b in old_blocks:
                        bbox = b.get("bbox", [])
                        if bbox and len(bbox) >= 4:
                            # block 完全在表格 bbox 内 → 移除
                            if (tx0 - 5 <= bbox[0] and bbox[2] <= tx1 + 5 and
                                    ty0 - 5 <= bbox[1] and bbox[3] <= ty1 + 5):
                                removed += 1
                                continue
                        kept.append(b)
                    page["para_blocks"] = kept + [new_block]
                    page["blocks"] = kept + [new_block]
                    print(f"  📋 第{pi+1}页补救表格: MinerU 漏检，从 PyMuPDF 重建"
                          f"（移除{removed}个 text block）", file=sys.stderr)

            # ── 用 PDF 大纲（书签）确定标题层级 ──
            # PDF 大纲是作者在 Word 里设置的大纲层级，是文档结构的权威定义。
            # MinerU 的 type=title 判断经常误判（如把"1.采购人信息"判为标题），
            # 用大纲坐标匹配来校验：只有在大纲里的 block 才是真标题。
            toc_matched = _match_blocks_with_toc(fdoc, pdf_info)
            fdoc.close()
            if line_h_found or widths_found or texts_found or toc_matched:
                print(f"  📐 排版信息提取: {line_h_found}个行高, "
                      f"{widths_found}个列宽, {texts_found}个单元格文本, "
                      f"{toc_matched}个大纲标题",
                      file=sys.stderr)
        except Exception as e:
            print(f"  ⚠️ 预扫描失败（回退默认）: {e}", file=sys.stderr)

    # ── 页面布局检测：从每页 block bbox 提取左边距、页宽等 ──
    # 替代所有硬编码值（82pt 左边距、595pt 页宽），支持不同版式的 PDF
    # 从 per_page_margins 提取全局边距回退（兼容 _detect_page_layout 的接口）
    if per_page_margins:
        all_t = [m["top"] for m in per_page_margins.values()]
        all_b = [m["bottom"] for m in per_page_margins.values()]
        all_l = [m["left"] for m in per_page_margins.values()]
        all_r = [m["right"] for m in per_page_margins.values()]
        margins = {"top": min(all_t), "bottom": min(all_b),
                   "left": min(all_l), "right": min(all_r)}
    else:
        margins = None
    _detect_page_layout(pdf_info, margins)

    # ── 行高兜底值（仅用于 bbox 无效的极端情况）──
    # 正常情况下每个 block 都能从 bbox高度/行数 直接取到精确行高，
    # 此兜底值仅在 bbox 数据缺失时使用。
    all_lh = []
    for page in pdf_info:
        for b in (page.get("para_blocks") or page.get("blocks") or []):
            lh = b.get("_line_height")
            if lh:
                all_lh.append(lh)
    doc_median_lh = 22.0  # 极端兜底
    if all_lh:
        all_lh.sort()
        doc_median_lh = all_lh[len(all_lh) // 2]

    for page in pdf_info:
        page_blocks = page.get("para_blocks") or page.get("blocks") or []
        for b in page_blocks:
            b["_doc_line_height"] = doc_median_lh

    # ── 段后间距：从同页相邻 block 的 Y 坐标差计算 ──
    for page in pdf_info:
        blocks = page.get("para_blocks") or page.get("blocks") or []
        for bi, block in enumerate(blocks):
            if bi + 1 < len(blocks):
                this_bbox = block.get("bbox", [])
                next_bbox = blocks[bi + 1].get("bbox", [])
                if this_bbox and next_bbox and len(this_bbox) >= 4 and len(next_bbox) >= 4:
                    gap = next_bbox[1] - this_bbox[3]  # 下一个 block 的 y0 - 当前 block 的 y1
                    if gap > 0:
                        block["_gap_after"] = gap

    block_count = 0
    prev_block_bottom = None  # 上一个 block 的 Y1 底部坐标

    from docx.enum.section import WD_SECTION

    for page in pdf_info:
        page_idx = page.get("page_idx", 0)
        page_size = page.get("page_size", [595, 842])
        page_height = page_size[1] if len(page_size) > 1 else 842
        blocks = page.get("para_blocks") or page.get("blocks") or []

        # 跳过空页（MinerU 无内容的 page_idx），避免产生空白 section
        has_content = any(
            any(s.get("content", "").strip()
                for ln in (b.get("lines") or []) for s in ln.get("spans", []))
            for b in blocks
            if (b.get("type") or "").lower() != "table"
        ) or any((b.get("type") or "").lower() == "table" for b in blocks)
        if not has_content:
            continue

        # ── 每页独立 section：section break 自带翻页效果 ──
        # 第一页用 section[0]（已在外层设置），后续页创建新 section
        # 用当前页的 page_size（支持横版页），而非首页全局值
        if page_idx > 0:
            from docx.enum.section import WD_ORIENT
            new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
            cur_w = page_size[0] if len(page_size) > 0 else 595
            cur_h = page_size[1] if len(page_size) > 1 else 842
            if cur_w > cur_h:  # 横版页
                new_sec.orientation = WD_ORIENT.LANDSCAPE
            else:
                new_sec.orientation = WD_ORIENT.PORTRAIT
            new_sec.page_width = Mm(cur_w * 0.3528)
            new_sec.page_height = Mm(cur_h * 0.3528)
            # 该页独立边距
            if per_page_margins and page_idx in per_page_margins:
                pm = per_page_margins[page_idx]
                new_sec.left_margin = Pt(pm["left"])
                new_sec.right_margin = Pt(pm["right"])
                new_sec.top_margin = Pt(pm["top"])
                new_sec.bottom_margin = Pt(pm["bottom"])
        prev_block_bottom = None  # 翻页后重置 Y 追踪

        for block in blocks:
            # 跨页续行 block：section break 已保证新页，只需重置 Y 追踪
            if block.get("_crosspage_continuation"):
                prev_block_bottom = None

            # 根据与上一个 block 的 Y 坐标差插入垂直留白
            bbox = block.get("bbox", [])
            if bbox and len(bbox) >= 4:
                y0 = bbox[1]
                if prev_block_bottom is not None:
                    gap = y0 - prev_block_bottom
                    # Y 差超过页高 6% 时，插入空段落还原留白
                    # 比例阈值替代硬编码 50pt，适应不同页面尺寸的 PDF
                    gap_threshold = page_height * 0.06
                    if gap > gap_threshold:
                        _add_vertical_gap(doc, gap)
                prev_block_bottom = bbox[3]

            btype = (block.get("type") or block.get("block_type") or "text").lower()
            builder = _BLOCK_BUILDERS.get(btype, _build_text)

            try:
                if btype in ("image", "picture"):
                    builder(doc, block, images_dir)
                else:
                    builder(doc, block)
                block_count += 1
            except Exception as e:
                print(f"  ⚠️ block 重建失败 (page={page_idx}, type={btype}): {e}",
                      file=sys.stderr)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    print(f"[4/4] DOCX 重建完成: {output_path}（{block_count} 个 block）",
          file=sys.stderr)
    return str(output_path.resolve())


# ═══════════════════════════════════════════════════════════════
#  回归检测：渲染后对比 PDF，微调表格行高消除页面偏移
# ═══════════════════════════════════════════════════════════════

def _render_docx_to_pdf(
    docx_path: str, out_dir: str,
) -> str | None:
    """用 LibreOffice headless 把 DOCX 转成 PDF，返回 PDF 路径。"""
    import subprocess, shutil, os
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        for p in (r"C:\Program Files\LibreOffice\program\soffice.exe",
                  r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"):
            if os.path.isfile(p):
                soffice = p
                break
    if not soffice:
        print("  ⚠️ 回归检测: 未找到 LibreOffice，跳过", file=sys.stderr)
        return None
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf",
             "--outdir", out_dir, docx_path],
            capture_output=True, timeout=120,
        )
        pdf_path = os.path.join(out_dir,
                                os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
        return pdf_path if os.path.isfile(pdf_path) else None
    except Exception:
        return None


def _normalize_page_text(txt: str) -> str:
    """归一化页面文本：去页码、去空白，取前40字符用于对比。"""
    import re
    lines = [l.strip() for l in txt.split("\n") if l.strip()]
    content = [l for l in lines if not re.match(r'^[—\-]?\s*\d{1,3}\s*[—\-]?$', l)]
    return re.sub(r'\s+', '', "".join(content))[:40]


def _check_page_alignment(
    docx_pdf_path: str, orig_pdf_path: str,
    pdf_info: list | None = None,
) -> tuple[list[int], int]:
    """
    对比 DOCX 渲染 PDF 和原始 PDF 的每页首行，返回偏移信息。

    如果传入 pdf_info，只检测含表格的页面（偏移主要由表格行高差异引起，
    正文区域已由分页符固定，检测正文页会因格式微差误报）。

    返回 (misaligned_pages, first_offset)
      - misaligned_pages: 不对齐的页码列表（1-based）
      - first_offset: 首个偏移页的偏移方向（+1=DOCX超前, -1=DOCX落后, 0=全对齐）
    """
    import fitz, re
    d = fitz.open(docx_pdf_path)
    o = fitz.open(orig_pdf_path)
    max_p = min(d.page_count, o.page_count)

    def _full_text(txt):
        """归一化整页文本（去页码去空白），用于内容搜索。"""
        lines = [l.strip() for l in txt.split("\n") if l.strip()]
        content = [l for l in lines if not re.match(r'^[—\-]?\s*\d{1,3}\s*[—\-]?$', l)]
        return re.sub(r'\s+', '', "".join(content))

    # 如果有 pdf_info，只检测含表格的页面
    table_pages = None
    if pdf_info:
        table_pages = set()
        for pi, page in enumerate(pdf_info):
            if not isinstance(page, dict):
                continue
            for b in (page.get("para_blocks") or page.get("blocks") or []):
                if (b.get("type") or "").lower() == "table":
                    table_pages.add(pi)
                    break

    misaligned = []
    first_offset = 0
    for pno in range(max_p):
        if table_pages is not None and pno not in table_pages:
            continue
        d_txt = _normalize_page_text(d[pno].get_text())
        o_txt = _normalize_page_text(o[pno].get_text())
        if d_txt == o_txt:
            continue
        misaligned.append(pno + 1)
        if first_offset == 0:
            # 双向检测偏移方向：
            # 方向A：DOCX首行在PDF哪页（DOCX超前=在上一页, 落后=在下一页）
            # 方向B：PDF首行在DOCX哪页（PDF首行在DOCX上一页=DOCX落后, 在下一页=DOCX超前）
            d_head = _full_text(d[pno].get_text())[:15]
            o_head = _full_text(o[pno].get_text())[:15]
            # 检查DOCX首行在PDF邻近页
            if d_head and pno > 0 and d_head in _full_text(o[pno - 1].get_text()):
                first_offset = 1   # DOCX 首行在 PDF 上一页 → DOCX 超前
            elif d_head and pno + 1 < o.page_count and d_head in _full_text(o[pno + 1].get_text()):
                first_offset = -1  # DOCX 首行在 PDF 下一页 → DOCX 落后
            # 反向检查：PDF首行在DOCX邻近页
            elif o_head and pno > 0 and o_head in _full_text(d[pno - 1].get_text()):
                first_offset = -1  # PDF 首行在 DOCX 上一页 → DOCX 落后（内容少了）
            elif o_head and pno + 1 < d.page_count and o_head in _full_text(d[pno + 1].get_text()):
                first_offset = 1   # PDF 首行在 DOCX 下一页 → DOCX 超前（内容多了）
            # 方向=0(未匹配)时继续找下一个不对齐页

    d.close()
    o.close()
    return misaligned, first_offset


def _auto_align_tables(
    merged_data: dict[str, Any], images_dir: str | None,
    output_path: str | Path, pdf_path: str,
    max_rounds: int = 5,
) -> str:
    """
    回归检测+行高微调：build → 渲染 → 检测偏移 → 调整表格行高补偿 → 重复。

    每轮：
    1. build_docx 生成 DOCX（带当前补偿值）
    2. LibreOffice 渲染成 PDF
    3. fitz 对比每页首行，找偏移
    4. 给偏移源的表格 block 加/减补偿值
    收敛条件：全对齐，或达到 max_rounds。
    """
    pdf_info = merged_data.get("pdf_info", merged_data if isinstance(merged_data, list) else [])
    output_path = Path(output_path)
    tmp_dir = str(output_path.parent / "_align_tmp")

    import os, shutil
    os.makedirs(tmp_dir, exist_ok=True)

    for round_num in range(1, max_rounds + 1):
        # 1. build
        result = build_docx(merged_data, images_dir, output_path, pdf_path=None)
        print(f"  🔄 回归检测第{round_num}轮: build完成", file=sys.stderr)

        # 2. 渲染
        os.makedirs(tmp_dir, exist_ok=True)
        # 把 output.docx 复制到 tmp（避免 LibreOffice 锁文件）
        tmp_docx = os.path.join(tmp_dir, "align_check.docx")
        import shutil
        shutil.copy2(str(output_path), tmp_docx)
        rendered_pdf = _render_docx_to_pdf(tmp_docx, tmp_dir)
        if not rendered_pdf:
            break  # 无法渲染，放弃

        # 3. 检测（只检测表格页的偏移）
        misaligned, offset = _check_page_alignment(rendered_pdf, pdf_path, pdf_info)
        if not misaligned:
            print(f"  ✅ 回归检测: 全部页对齐", file=sys.stderr)
            break

        if offset == 0:
            # 有不对齐页但无法判断方向（如目录点号差异），视为收敛
            print(f"  ✅ 回归检测: {len(misaligned)}页有格式差异但无内容偏移，收敛",
                  file=sys.stderr)
            break

        print(f"  📍 偏移方向={offset:+d}, 不对齐共{len(misaligned)}页: {misaligned[:10]}",
              file=sys.stderr)

        # 4. 给偏移页附近的表格 block 加补偿
        # DOCX超前(offset=+1) → 表格矮了 → 需要加高行高 → 补偿 +=
        # DOCX落后(offset=-1) → 表格高了 → 需要减高行高 → 补偿 -=
        # 只调第一个偏移页的表格（避免过调），每轮加3pt/行（快速收敛）
        adjusted = _adjust_table_compensation(pdf_info, misaligned[0], offset)
        if not adjusted:
            print(f"  ⚠️ 偏移页附近未找到可调整的表格，停止", file=sys.stderr)
            break

    # 最终 build（确保用最新补偿值）
    result = build_docx(merged_data, images_dir, output_path, pdf_path=pdf_path)
    return result


def _adjust_table_compensation(
    pdf_info: list, bad_page_1based: int, offset: int,
) -> bool:
    """
    给偏移源的表格 block 调整行高补偿值。

    bad_page_1based: 首个有明确方向的偏移页（1-based）
    offset: +1=DOCX超前(需加高), -1=DOCX落后(需减高)

    策略：从偏移页向前搜索最近的表格 block（偏移由前方表格矮/高引起），
    给它们的行高补偿叠加微调。每轮微调量固定为2pt/行（保守，避免过冲）。
    """
    target_pi = bad_page_1based - 1  # 0-based page_idx
    # 从偏移页向前搜索最近的表格（最多搜5页）
    check_pages = list(range(max(0, target_pi - 5), target_pi + 1))
    adjusted = False

    for pi in check_pages:
        if pi < 0 or pi >= len(pdf_info):
            continue
        page = pdf_info[pi]
        if not isinstance(page, dict):
            continue
        blocks = page.get("para_blocks") or page.get("blocks") or []
        for b in blocks:
            if (b.get("type") or "").lower() != "table":
                continue
            if b.get("_table_merged"):
                continue
            # 补偿值不依赖 _table_row_heights（该字段在 build 内部生成，不存 JSON）
            # 只要 block 是 table 类型且有 bbox（build 时会提取行高）就加补偿
            # 计算补偿方向：
            # offset=+1 (DOCX超前,内容多了) → 表格矮了 → 加高 → comp +=
            # offset=-1 (DOCX落后,内容少了) → 表格高了 → 减高 → comp -=
            sign = 1 if offset > 0 else -1
            # 累积偏移93pt分布在33行表格中 ≈ 2.8pt/行，3轮迭代每轮约1pt/行
            delta = sign * 1.0
            old_comp = b.get("_row_height_compensation", 0)
            b["_row_height_compensation"] = old_comp + delta
            adjusted = True
            print(f"    → pi={pi} 表格补偿: {old_comp:+.1f} → {b['_row_height_compensation']:+.1f}pt/行",
                  file=sys.stderr)
    return adjusted


def _add_vertical_gap(doc, gap_pt: float) -> None:
    """
    插入垂直留白（用空段落的段前距实现）。
    gap_pt: 需要的间距（PDF 点，1pt ≈ 0.35mm）。
    """
    from docx.shared import Pt
    # gap_pt 是 PDF 坐标的间距，直接转为磅值（1pt = 1磅）
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = Pt(1)  # 最小行高
    # 用段前距实现间距
    para.paragraph_format.space_before = Pt(gap_pt)


def _add_page_break(doc) -> None:
    """插入分页符（段落行高设为最小，避免表格填满页时分页符段落溢出产生空页）。"""
    from docx.enum.text import WD_BREAK
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(1)  # 最小行高，避免空段落占整行空间
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)


# ═══════════════════════════════════════════════════════════════
#  CLI（可独立运行测试）
# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import argparse
    import json

    parser = argparse.ArgumentParser(
        description="python-docx DOCX 重建"
    )
    parser.add_argument("merged", help="align.py 输出的 merged.json 路径")
    parser.add_argument("-o", "--output", default="./output/重建结果.docx",
                        help="输出 docx 路径")
    parser.add_argument("--images-dir", default=None,
                        help="图片目录路径")

    args = parser.parse_args()

    merged_data = json.loads(Path(args.merged).read_text(encoding="utf-8"))
    result = build_docx(merged_data, args.images_dir, args.output)

    print(json.dumps({"output": result}, ensure_ascii=False, indent=2))
